"""Generate DigiLog Functional Documentation as a formatted Word document."""
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


OUTPUT = Path(__file__).resolve().parent.parent / "docs" / "DigiLog_Functional_Documentation.docx"


def set_cell_shading(cell, hex_color: str):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), hex_color)
    cell._tc.get_or_add_tcPr().append(shading)


def add_table(doc, headers, rows, header_fill="1F4E79"):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.size = Pt(10)
        set_cell_shading(hdr[i], header_fill)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            table.rows[ri + 1].cells[ci].text = str(val)
            for p in table.rows[ri + 1].cells[ci].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
    doc.add_paragraph()
    return table


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")
        for run in p.runs:
            run.font.size = Pt(11)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(item, style="List Number")
        for run in p.runs:
            run.font.size = Pt(11)


def build_document():
    doc = Document()

    # Default font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # ── Cover page ──
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("DigiLog")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(31, 78, 121)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = sub.add_run("Functional Documentation")
    r2.bold = True
    r2.font.size = Pt(18)
    r2.font.color.rgb = RGBColor(68, 84, 106)

    doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run("Zuari Industries — Digital Operations Platform\n").font.size = Pt(12)
    meta.add_run(f"Document Version: 1.0\n").font.size = Pt(11)
    meta.add_run(f"Date: {date.today().strftime('%d %B %Y')}\n").font.size = Pt(11)
    meta.add_run(
        "Audience: Plant Operators, Department Users, Managers, Business Stakeholders, Auditors, Process Owners"
    ).font.size = Pt(11)

    doc.add_page_break()

    # ── Document control ──
    doc.add_heading("Document Control", level=1)
    add_table(
        doc,
        ["Item", "Details"],
        [
            ["Document Title", "DigiLog Functional Documentation"],
            ["System Name", "DigiLog — Digital Logbook Platform"],
            ["Organization", "Zuari Industries"],
            ["Purpose", "Describe business functionality and user workflows for non-technical stakeholders"],
            ["Classification", "Internal — Business Use"],
        ],
    )

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════
    # 1. APPLICATION OVERVIEW
    # ═══════════════════════════════════════════════════════════════
    doc.add_heading("1. Application Overview", level=1)

    doc.add_heading("1.1 What is DigiLog?", level=2)
    doc.add_paragraph(
        "DigiLog is Zuari Industries' centralized digital operations platform. It replaces paper-based "
        "logbooks, equipment history cards, and manual reporting with a secure web application accessible "
        "from plant workstations. Operators and department staff capture shift readings, stoppages, "
        "laboratory results, power and steam data, distillery operations, safety incidents, and production "
        "logs in structured digital forms. Managers and leadership access consolidated dashboards for "
        "operational visibility and analytics."
    )

    doc.add_heading("1.2 Why DigiLog Exists", level=2)
    add_bullets(
        doc,
        [
            "Paper logbooks are difficult to search, consolidate, and audit.",
            "Operational data is scattered across departments with inconsistent formats.",
            "Leadership needs timely visibility into mill, lab, power, distillery, and production performance.",
            "Regulatory and internal audit requirements demand traceable, retrievable records.",
            "Equipment maintenance history must be preserved in a single, accessible repository.",
        ],
    )

    doc.add_heading("1.3 Business Problems Solved", level=2)
    add_table(
        doc,
        ["Problem", "DigiLog Solution"],
        [
            ["Manual paper logbooks", "Digital forms with validation and permanent storage"],
            ["Delayed reporting", "Real-time data entry and BI dashboards"],
            ["Lost or incomplete records", "Centralized database with view and CSV export"],
            ["Equipment knowledge in files/folders", "Digital equipment history cards with specs, OEM schedule, and maintenance timeline"],
            ["Access control gaps", "Role-based access with admin-assigned module mappings"],
            ["Audit preparation effort", "Searchable historical records and exportable data"],
        ],
    )

    doc.add_heading("1.4 Benefits of Digitization", level=2)
    add_bullets(
        doc,
        [
            "Single source of truth for operational registers across the plant.",
            "Standardized data capture with mandatory fields and validation rules.",
            "Faster retrieval of historical readings for troubleshooting and audits.",
            "Analytics dashboards for distillery and milling operations.",
            "Reduced physical storage and handling of paper registers.",
            "Secure authentication including corporate SSO (Microsoft / Google).",
        ],
    )

    doc.add_heading("1.5 Departments Using the System", level=2)
    add_table(
        doc,
        ["Department / Area", "Primary Modules"],
        [
            ["Milling", "Mill Logbook (temperatures, shredder, lube, stoppages), Milling BI Cockpit"],
            ["Laboratory", "DS, RS, Operations, Special Analysis, Syrup, Stoppage logbooks"],
            ["Power House", "Power, Steam, Stoppage logbooks; Power Plant Equipment History"],
            ["Distillery", "Distillery Operations logbook; Distillery Analytics dashboard"],
            ["Mill House Maintenance", "Mill House Equipment History Cards"],
            ["Engineering (Electrical / Instrument / Mechanical)", "Power Plant Equipment History (legacy and new hierarchy)"],
            ["EHS", "Near Miss / Incident reports; Water dashboards (GWA, ETP, CPU)"],
            ["Production", "Shift chemist, centrifugal, pan, decanter, clarification logbooks"],
            ["Administration / IT", "User management, access mapping, system configuration"],
            ["Leadership / Analytics", "BI Control Tower dashboards"],
        ],
    )

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════
    # 2. BUSINESS OBJECTIVES
    # ═══════════════════════════════════════════════════════════════
    doc.add_heading("2. Business Objectives", level=1)
    objectives = [
        ("Operational Goals", "Enable consistent, timely capture of shift and daily operational data across all plant areas."),
        ("Process Standardization", "Replace varied paper formats with uniform digital forms aligned to plant SOPs."),
        ("Reduction of Manual Logbooks", "Phase out paper registers for mill, lab, power, distillery, EHS, and production areas."),
        ("Traceability", "Maintain a permanent, timestamped record of each submitted entry in a central database."),
        ("Audit Readiness", "Support internal and external audits with searchable history and CSV export."),
        ("Reporting Improvements", "Provide BI dashboards for distillery operations and milling division performance."),
        ("Operational Visibility", "Give managers and leadership near-real-time insight through dashboards and consolidated data views."),
    ]
    for title, desc in objectives:
        p = doc.add_paragraph()
        p.add_run(f"{title}: ").bold = True
        p.add_run(desc)

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════
    # 3. USER ROLES
    # ═══════════════════════════════════════════════════════════════
    doc.add_heading("3. User Roles", level=1)
    doc.add_paragraph(
        "DigiLog implements two system roles. Business titles such as Supervisor, Reviewer, or Auditor "
        "map to these roles plus admin-assigned module access — the platform does not define separate "
        "supervisor or approver login types."
    )

    doc.add_heading("3.1 Employee (Operator / Department User)", level=2)
    doc.add_paragraph("Typical users: shift operators, lab chemists, power house staff, EHS officers, production personnel.")
    doc.add_heading("Responsibilities", level=3)
    add_bullets(
        doc,
        [
            "Enter operational data into assigned digital logbooks.",
            "Review entries in the confirmation modal before committing.",
            "View historical records and export CSV for assigned forms.",
            "Maintain equipment history where authorized.",
            "Access BI dashboards when mapped by admin.",
        ],
    )
    doc.add_heading("Accessible Modules", level=3)
    doc.add_paragraph("Only modules and forms assigned by the administrator. Typical assignments include Mill Logbook, Lab Logbook, Power Logbook, Distillery, EHS, Production, Equipment History, and BI Control Tower.")
    doc.add_heading("Actions Allowed", level=3)
    add_bullets(doc, ["Submit log entries", "View data and download CSV", "Update equipment specs/history (equipment modules)", "Upload data files (if granted)", "View BI dashboards (if mapped)"])

    doc.add_heading("3.2 Administrator", level=2)
    doc.add_paragraph("Typical users: IT administrators, system owners, HR/operations coordinators managing access.")
    doc.add_heading("Responsibilities", level=3)
    add_bullets(
        doc,
        [
            "Create and manage employee accounts.",
            "Assign module and form access (mappings).",
            "Send account activation emails with temporary passwords.",
            "Configure BI dashboard settings.",
            "Grant data upload access to employees.",
            "Assign reporting manager (informational field).",
        ],
    )
    doc.add_heading("Accessible Modules", level=3)
    doc.add_paragraph("Admin portal only: Employee Management. Admins use a separate login page and cannot access the employee portal with their admin account.")
    doc.add_heading("Actions Allowed", level=3)
    add_bullets(doc, ["Full user CRUD", "Form and dashboard mapping", "Activate/deactivate users", "Bulk email activation", "BI settings (e.g. third season comparison toggle)", "Data upload access management"])

    doc.add_heading("3.3 Business Role Mapping", level=2)
    add_table(
        doc,
        ["Business Title", "DigiLog Role", "Typical Access"],
        [
            ["Plant Operator", "Employee", "Assigned logbooks for their shift/area"],
            ["Lab Chemist", "Employee", "Lab logbook forms"],
            ["Supervisor / Section Head", "Employee", "Same as operator; signature fields captured as text in forms — no separate approval queue"],
            ["Department Head", "Employee or Admin", "BI dashboards; admin if managing users"],
            ["Auditor", "Employee (read-only by policy)", "View Data + CSV export on assigned forms; no dedicated auditor role"],
            ["Process Owner", "Admin", "User mapping and module configuration"],
            ["EHS Officer", "Employee", "EHS forms and water dashboards"],
        ],
    )

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════
    # 4. FUNCTIONAL MODULES
    # ═══════════════════════════════════════════════════════════════
    doc.add_heading("4. Functional Modules", level=1)

    modules = [
        {
            "name": "4.1 Mill Logbook",
            "purpose": "Capture milling division operational readings and stoppages.",
            "use_case": "Shift operators record equipment temperatures, shredder/OTG readings, lube pressures, roller temperatures, and mill stoppages.",
            "features": ["Equipment Temperature (mill_logbook1)", "Shredder and OTG (mill_logbook2)", "Lube Pressure and Roller Temp (mill_logbook3)", "Mill Stoppages (mill_stoppages)"],
            "workflow": "Select form → Enter Date/Shift/Time and readings → Review modal → Confirm & Commit.",
            "validations": "Date and Shift mandatory; duplicate entries blocked by Date+Shift+Time (or Date+start/end for stoppages).",
            "reports": "View Data modal; CSV export; feeds Milling BI Cockpit.",
        },
        {
            "name": "4.2 Lab Logbook",
            "purpose": "Record laboratory analysis results for sugar manufacturing process control.",
            "use_case": "Lab staff enter pol, brix, and related measurements by shift and sampling time.",
            "features": ["DS Logbook", "RS Logbook", "Operations Logbook", "Special Analysis Logbook", "Syrup Logbook", "Stoppage Logbook"],
            "workflow": "Standard submit with review modal; uniqueness by Date+Shift+Sampling time (or Date+Shift for syrup).",
            "validations": "Required date/shift fields; duplicate prevention on single submit.",
            "reports": "View Data; CSV export.",
        },
        {
            "name": "4.3 Power Logbook",
            "purpose": "Document power generation, steam balance, and power house stoppages.",
            "use_case": "Power house operators log hourly/daily power and steam parameters and stoppage events.",
            "features": ["Power Details (ph_power)", "Steam Details (ph_steam)", "Stoppage Details (ph_stoppage)"],
            "workflow": "Date + Time based entry; review and commit.",
            "validations": "Duplicate blocked per Date+Time.",
            "reports": "View Data; CSV export.",
        },
        {
            "name": "4.4 Distillery Operations",
            "purpose": "Daily snapshot of distillery production, efficiency, and storage metrics.",
            "use_case": "Distillery staff record operation mode, ethanol production, molasses consumption, efficiencies, and storage levels.",
            "features": ["Single daily form (distillery_ops)", "Auto-calculated fields in database (FS%, total molasses in store)"],
            "workflow": "One record per date; review and commit.",
            "validations": "One entry per calendar date.",
            "reports": "View Data; CSV export; Distillery Analytics BI dashboard.",
        },
        {
            "name": "4.5 Mill House Equipment History",
            "purpose": "Digital equipment life history cards for mill house assets.",
            "use_case": "Maintenance teams maintain specifications, OEM schedules, and maintenance history with photos.",
            "features": ["Equipment list with search", "Specifications hub", "OEM maintenance schedule", "Maintenance history timeline with before/after images"],
            "workflow": "Browse equipment → View/edit specs, schedule, history → Add history entries.",
            "validations": "Image upload for equipment photos; structured history fields.",
            "reports": "In-module history view; Excel export for specs and schedule.",
        },
        {
            "name": "4.6 Power Plant Equipment History (Legacy)",
            "purpose": "Equipment cards organized by engineering discipline (Electrical, Instrument, etc.).",
            "use_case": "Engineering teams maintain 30MW-era equipment records by department.",
            "features": ["Department-based navigation", "Same specs/schedule/history structure as mill house"],
            "workflow": "Select department → Select equipment → Manage card content.",
            "validations": "Standard equipment history validations.",
            "reports": "History timeline; Excel exports.",
        },
        {
            "name": "4.7 Power Plant Equipment History (New Hierarchy)",
            "purpose": "Plant-wide equipment hierarchy (150TPH, 70TPH, WTP, etc.) with manageable tree structure.",
            "use_case": "Engineering teams navigate boiler/turbine/WTP hierarchy; admins can add/rename/delete nodes (with protection for seeded nodes).",
            "features": ["Hierarchy explorer", "Discipline-based detail pages", "Sub-group history scoping", "Hierarchy management panel (admin-capable users)"],
            "workflow": "Tree navigation → Equipment detail → History/specs/schedule.",
            "validations": "Protected seeded hierarchy nodes cannot be edited.",
            "reports": "Equipment history; Excel exports.",
        },
        {
            "name": "4.8 EHS — Environment, Health & Safety",
            "purpose": "Incident reporting and water management dashboards.",
            "use_case": "EHS staff report near misses/incidents and daily water abstraction, ETP, and CPU metrics.",
            "features": ["Near Miss / Incident / Accident Report", "Water Dashboard — Ground Water Abstraction", "Water Dashboard — ETP Working", "Water Dashboard — CPU Water Recycle"],
            "workflow": "Standard forms with HOD signature fields captured as text at entry time.",
            "validations": "Required fields on near miss; pH values must be 0–14 on water forms.",
            "reports": "View Data; CSV export.",
        },
        {
            "name": "4.9 Production",
            "purpose": "Production department shift and machine logbooks.",
            "use_case": "Production chemists and operators log shift activities, centrifugal stoppages, pan operations, decanter hourly readings, and clarification data.",
            "features": ["Shift Chemist Job Log Book", "A-Centrifugal Machine Stoppage", "Pan Log Book (batch)", "Decanter Log Book (batch, 24 slots)", "Clarification Log Book (batch)"],
            "workflow": "Single or batch submit with review modal.",
            "validations": "Daily uniqueness by date; batch forms allow multiple rows per session.",
            "reports": "View Data; CSV export.",
        },
        {
            "name": "4.10 BI Control Tower",
            "purpose": "Executive and operational analytics dashboards.",
            "use_case": "Managers analyze distillery and milling performance with KPIs, trends, and comparisons.",
            "features": ["Distillery Operations Analytics", "Milling Division Cockpit (Outage, Equipment Temp, Lube/Roller tabs)"],
            "workflow": "Open dashboard → Select date range and filters → View KPIs and charts.",
            "validations": "Access controlled by dashboard mapping.",
            "reports": "In-dashboard raw data tables; historical source via form CSV export.",
        },
        {
            "name": "4.11 Data Ingestion Center",
            "purpose": "Controlled upload of reference data files (CSV/Excel).",
            "use_case": "Authorized users upload mill mapping files that feed BI thermal and lube reports.",
            "features": ["Category tagging", "File list with uploader audit", "Download and delete own uploads", "Auto-sync for recognized mill mapping filenames"],
            "workflow": "Check access → Upload file with category → File appears in registry.",
            "validations": "Category 3–200 characters; CSV/XLS/XLSX only; admin-granted access required.",
            "reports": "Upload history list with uploader name and timestamp.",
        },
        {
            "name": "4.12 Administration",
            "purpose": "System governance and user lifecycle management.",
            "use_case": "Administrators provision users, control access, and configure BI options.",
            "features": ["Employee table with search", "Create/edit/deactivate/delete users", "Form mapping modal", "Dashboard mapping modal", "Send activation email", "Manager assignment", "BI settings", "Data upload access"],
            "workflow": "Admin login → Employee Management → Configure users and mappings.",
            "validations": "Email uniqueness; SSO users cannot receive password emails.",
            "reports": "Employee list export via UI table.",
        },
    ]

    for m in modules:
        doc.add_heading(m["name"], level=2)
        for label, key in [
            ("Purpose", "purpose"),
            ("Business Use Case", "use_case"),
            ("Workflow", "workflow"),
            ("Validations", "validations"),
            ("Reports Generated", "reports"),
        ]:
            p = doc.add_paragraph()
            p.add_run(f"{label}: ").bold = True
            p.add_run(m[key])
        p = doc.add_paragraph()
        p.add_run("Key Features:").bold = True
        add_bullets(doc, m["features"])

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════
    # 5. DAILY WORKFLOW
    # ═══════════════════════════════════════════════════════════════
    doc.add_heading("5. Daily Workflow", level=1)
    doc.add_paragraph("The following describes a typical operational day for an employee user.")

    doc.add_heading("5.1 Step-by-Step User Journey", level=2)
    steps = [
        "Sign in via the DigiLog home page using email/password, Microsoft, or Google (account must be pre-created by admin).",
        "Land on the Homepage showing assigned destination cards: Forms Hub and/or BI Control Tower.",
        "Open Forms Hub to see assigned operational modules (Mill, Lab, Power, etc.).",
        "Select an application (e.g., Mill Logbook) to see available forms.",
        "Open the required form for the current shift (e.g., Equipment Temperature).",
        "Enter Date, Shift (A/B/C/G where applicable), Time, and all required readings.",
        "Click Submit — the Review modal displays a summary of entered values.",
        "Read the certification notice and click Confirm & Commit to permanently save the entry.",
        "Repeat for additional forms during the shift.",
        "Optionally open BI Control Tower to review operational KPIs and trends.",
        "At end of shift, ensure all mandatory logbooks for the area are completed.",
    ]
    add_numbered(doc, steps)

    doc.add_heading("5.2 Shift Handling", level=2)
    doc.add_paragraph(
        "Most mill and lab forms require a Shift field (A, B, C, or G). Combined with Date and Time (or sampling time), "
        "this ensures each shift's readings are recorded separately. Stoppage forms use start and end timestamps instead."
    )

    doc.add_heading("5.3 Corrections", level=2)
    doc.add_paragraph(
        "Currently, committed logbook entries cannot be edited or deleted through the user interface. If an incorrect "
        "entry was submitted, operational procedure should define whether a corrective entry is made or IT/admin "
        "intervention is required. The system displays a duplicate-prevention message referencing editing, but "
        "in-app edit is not yet available."
    )

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════
    # 6. LOG BOOK FUNCTIONALITY
    # ═══════════════════════════════════════════════════════════════
    doc.add_heading("6. Log Book Functionality", level=1)

    doc.add_heading("6.1 Types of Logs Maintained", level=2)
    add_table(
        doc,
        ["Category", "Forms", "Frequency Pattern"],
        [
            ["Mill", "Equipment Temp, Shredder/OTG, Lube/Roller, Stoppages", "Per shift / per stoppage event"],
            ["Lab", "DS, RS, Ops, SA, Syrup, Stoppage", "Per shift / sampling time"],
            ["Power", "Power, Steam, Stoppage", "Per time slot / event"],
            ["Distillery", "Daily operations snapshot", "One per calendar date"],
            ["EHS", "Near miss, Water GWA/ETP/CPU", "Per event / daily"],
            ["Production", "Chemist, Centrifugal, Pan, Decanter, Clarification", "Per shift / hourly batch"],
        ],
    )

    doc.add_heading("6.2 Data Captured", level=2)
    doc.add_paragraph(
        "Each form captures domain-specific measurements: temperatures, pressures, pol/brix values, power and steam "
        "flows, stoppage durations and reasons, production quantities, water quality parameters, and incident details. "
        "All entries receive a system timestamp at commit time."
    )

    doc.add_heading("6.3 Mandatory Fields", level=2)
    add_bullets(
        doc,
        [
            "Date — required on virtually all forms.",
            "Shift — required on mill, lab, and many production forms.",
            "Time or Sampling Time — required where the form pattern demands it.",
            "Form-specific required fields (e.g., person name and severity on EHS Near Miss).",
        ],
    )

    doc.add_heading("6.4 Validation Checks", level=2)
    add_bullets(
        doc,
        [
            "Client-side required field validation before submit.",
            "Server-side duplicate detection (single submit) based on date/shift/time keys.",
            "pH range validation (0–14) on EHS water forms.",
            "Email format validation at login.",
            "Duplicate operation blocked with user-friendly message.",
        ],
    )

    doc.add_heading("6.5 Timestamps", level=2)
    doc.add_paragraph(
        "Each committed record receives an automatic system timestamp. This supports chronological sorting in "
        "View Data and BI analytics. Note: the submitting user's identity is not currently stored on logbook rows."
    )

    doc.add_heading("6.6 Attachments and Images", level=2)
    doc.add_paragraph(
        "Logbook forms do not support general file attachments. Equipment history modules support before/after "
        "maintenance photos. User profile photos (avatars) are supported separately in account settings."
    )

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════
    # 7. APPROVAL & REVIEW WORKFLOW
    # ═══════════════════════════════════════════════════════════════
    doc.add_heading("7. Approval & Review Workflow", level=1)

    doc.add_paragraph(
        "IMPORTANT: DigiLog does not implement a formal multi-level approval workflow. There are no pending, "
        "approved, or rejected states, no reviewer queues, and no manager sign-off gates before data is saved."
    )

    doc.add_heading("7.1 Pre-Submit Review (Self-Certification)", level=2)
    add_numbered(
        doc,
        [
            "User completes the form and clicks Submit.",
            "Review modal opens showing all entered values in a read-only summary.",
            "User reads the certification statement confirming accuracy.",
            "User clicks Confirm & Commit — record is permanently saved.",
        ],
    )

    doc.add_heading("7.2 Signature and HOD Fields", level=2)
    doc.add_paragraph(
        "Some forms (EHS Near Miss, Production logbooks) include fields for HOD comments, signatures, or section head "
        "names. These are captured as text during data entry by the same user — they do not trigger a separate "
        "approval step or notification to the named person."
    )

    doc.add_heading("7.3 Rejection and Resubmission", level=2)
    doc.add_paragraph(
        "Not applicable — there is no rejection workflow. Duplicate entries are blocked at submit time; users must "
        "change date/shift/time to submit a corrected reading as a new entry (subject to duplicate rules)."
    )

    doc.add_heading("7.4 Escalation", level=2)
    doc.add_paragraph(
        "No automated escalation exists. The manager field on employee records is informational only and is not "
        "linked to workflows or notifications."
    )

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════
    # 8. DASHBOARD & REPORTS
    # ═══════════════════════════════════════════════════════════════
    doc.add_heading("8. Dashboard & Reports", level=1)

    doc.add_heading("8.1 Homepage Dashboard", level=2)
    doc.add_paragraph(
        "After login, employees see up to two destination cards: Forms Hub (operational modules) and BI Control Tower "
        "(analytics). Card visibility is controlled entirely by administrator mappings."
    )

    doc.add_heading("8.2 Distillery Operations Analytics", level=2)
    add_bullets(
        doc,
        [
            "KPI cards: production volumes, efficiencies, recovery, storage levels.",
            "Period comparisons: prior period, season-to-season (admin can enable third season comparison).",
            "Operation mode filter (B Heavy, C Heavy, Syrup, Mixed).",
            "Date range presets (MTD, last 30 days, season, custom).",
            "Trend charts for production, efficiency, and storage metrics.",
            "Raw data table within the dashboard.",
        ],
    )

    doc.add_heading("8.3 Milling Division Cockpit", level=2)
    add_bullets(
        doc,
        [
            "Tab 1 — Mill Outage: stoppage analytics by section (Cane, Mills, Boilers, etc.).",
            "Tab 2 — Equipment Temperature: thermal reports from mill logbook data.",
            "Tab 3 — Lube & Roller Temp: lube pressure and roller temperature trends.",
            "Section multi-select filter and date range controls.",
            "KPI cards with sparkline trends.",
        ],
    )

    doc.add_heading("8.4 Form-Level Reports", level=2)
    add_bullets(
        doc,
        [
            "View Data — paginated modal showing historical submitted records (20 rows per page).",
            "Download CSV — export up to 10,000 records for offline analysis in Excel.",
        ],
    )

    doc.add_heading("8.5 Equipment Reports", level=2)
    doc.add_paragraph("Equipment modules support Excel export of specifications and OEM maintenance schedules from the equipment detail pages.")

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════
    # 9. NOTIFICATIONS & ALERTS
    # ═══════════════════════════════════════════════════════════════
    doc.add_heading("9. Notifications & Alerts", level=1)
    add_table(
        doc,
        ["Notification Type", "Available?", "Description"],
        [
            ["Account activation email", "Yes", "Admin sends email with temporary password and login link (local accounts only)"],
            ["Bulk activation email", "Yes", "Admin can send activation emails to multiple users at once"],
            ["In-app success/error toasts", "Yes", "Brief messages after submit, login, or errors"],
            ["Pending approval reminders", "No", "No approval workflow exists"],
            ["Missed entry alerts", "No", "System does not remind users of unsubmitted logbooks"],
            ["Operational incident alerts", "No", "EHS submissions do not trigger email/SMS to supervisors"],
            ["Threshold breach alerts", "No", "BI dashboards are passive — no proactive alerting"],
            ["In-app notification center", "No", "No persistent notification inbox"],
        ],
    )

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════
    # 10. SEARCH & AUDIT FEATURES
    # ═══════════════════════════════════════════════════════════════
    doc.add_heading("10. Search & Audit Features", level=1)

    doc.add_heading("10.1 Historical Search", level=2)
    add_bullets(
        doc,
        [
            "Form View Data modal — browse paginated history sorted by date descending.",
            "Equipment list search — find mill house equipment by name or equipment number.",
            "Admin employee search — filter by name, email, or department.",
            "BI dashboard date range filters — analyze specific periods.",
        ],
    )

    doc.add_heading("10.2 Audit Tracking", level=2)
    add_bullets(
        doc,
        [
            "System timestamp on every committed logbook record.",
            "Data upload audit: uploader name, email, category, file size, upload timestamp.",
            "Equipment history: created/updated timestamps on maintenance entries.",
            "Limitation: logbook entries do not record which user submitted the data.",
            "Limitation: no change audit trail — records cannot be edited after commit.",
        ],
    )

    doc.add_heading("10.3 Traceability", level=2)
    doc.add_paragraph(
        "Traceability is supported through immutable committed records, timestamps, and CSV export for external "
        "audit packages. Formal who-submitted attribution and edit history are recommended future enhancements."
    )

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════
    # 11. BUSINESS RULES
    # ═══════════════════════════════════════════════════════════════
    doc.add_heading("11. Business Rules", level=1)

    rules = [
        ("Access Control", "Employees see only admin-assigned modules and forms. Empty form mapping = all forms in the assigned app."),
        ("Portal Separation", "Admins must use the admin login portal; employees use the home portal. Cross-portal login is blocked."),
        ("SSO Pre-Registration", "Microsoft and Google users must be created by admin before first SSO login."),
        ("Mandatory Submissions", "Operational policy defines which logbooks must be completed per shift — enforced by procedure, not by system reminders."),
        ("Duplicate Prevention", "Single-submit forms block duplicate entries for the same date/shift/time combination."),
        ("Batch Submit", "Production pan/decanter/clarification forms allow multiple rows in one session; duplicate check is not applied to batch."),
        ("Data Permanence", "Committed logbook entries cannot be edited or deleted via the UI."),
        ("pH Validation", "EHS water form pH values must be between 0 and 14."),
        ("Data Upload Access", "Only admin-granted employees (and admins) can upload files to the Data Ingestion Center."),
        ("Equipment Access", "Any authenticated user can update equipment history — not restricted by app mapping."),
        ("Hierarchy Protection", "Seeded power plant hierarchy nodes (150TPH, 70TPH, WTP areas) cannot be renamed or deleted."),
        ("BI Data Window", "BI dashboards load data for the last 365 days by default to ensure performance."),
    ]
    for title, desc in rules:
        p = doc.add_paragraph()
        p.add_run(f"{title}: ").bold = True
        p.add_run(desc)

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════
    # 12. BENEFITS TO ORGANIZATION
    # ═══════════════════════════════════════════════════════════════
    doc.add_heading("12. Benefits to Organization", level=1)
    benefits = [
        "Improved accountability through standardized digital capture and self-certification at commit.",
        "Centralized operational data replacing scattered paper registers.",
        "Easier audits with searchable history and CSV export capabilities.",
        "Operational transparency via BI dashboards for distillery and milling performance.",
        "Reduced paperwork, printing, and physical storage requirements.",
        "Faster reporting — leadership can access dashboards without waiting for manual consolidation.",
        "Preserved equipment knowledge in digital history cards accessible across shifts.",
        "Controlled access ensuring users see only relevant modules for their role.",
        "Foundation for future enhancements (approvals, alerts, mobile access, submitter attribution).",
    ]
    add_bullets(doc, benefits)

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════
    # 13. TYPICAL USER SCENARIOS
    # ═══════════════════════════════════════════════════════════════
    doc.add_heading("13. Typical User Scenarios", level=1)

    scenarios = [
        (
            "13.1 Operator Submitting Shift Log",
            "Rajesh, a mill shift operator, logs into DigiLog at the start of Shift B. He opens Forms Hub → "
            "Mill Logbook → Equipment Temperature. He selects today's date, Shift B, and the current time, then "
            "enters motor and bearing temperatures for each mill section. He submits, reviews the summary in the "
            "confirmation modal, and clicks Confirm & Commit. He repeats for Shredder/OTG and Lube Pressure forms. "
            "His entries are immediately available for the Milling BI Cockpit."
        ),
        (
            "13.2 Lab Chemist Recording Analysis",
            "Priya, a lab chemist, accesses the Lab Logbook module and opens the DS Logbook. She enters pol and brix "
            "values for each sampling point with the correct shift and sampling time. The system prevents her from "
            "accidentally duplicating an entry for the same date, shift, and sampling time. She exports the week's "
            "data to CSV for her department review meeting."
        ),
        (
            "13.3 Supervisor Reviewing Entries",
            "Vikram, a section supervisor, opens the Mill Stoppages form and clicks View Data to review stoppages "
            "recorded during the past week. He filters by reviewing entries in the modal and cross-references the "
            "Milling Division Cockpit outage tab for section-wise downtime analysis. Note: Vikram does not approve "
            "entries in a workflow — he reviews data that operators have already committed."
        ),
        (
            "13.4 Admin Managing Users",
            "Anita, the system administrator, logs into the admin portal. She creates a new employee account for a "
            "joining operator, assigns Mill Logbook and Forms Hub mapping, sends an activation email with a temporary "
            "password, and verifies the employee appears in the mapping modal with the correct forms selected."
        ),
        (
            "13.5 Auditor Checking Historical Logs",
            "An external auditor requests distillery production records for the previous season. The process owner "
            "exports CSV data from the Distillery Operations form and provides BI dashboard screenshots. The auditor "
            "cross-references timestamps and daily entries. Limitation: the auditor cannot determine individual "
            "submitters from the data alone."
        ),
        (
            "13.6 Maintenance Engineer Updating Equipment History",
            "Suresh, a maintenance engineer, navigates to Power Plant Equipment History (New), finds a boiler feed "
            "pump in the hierarchy, and adds a maintenance history entry with observation, action taken, cost, and "
            "before/after photos from a recent overhaul."
        ),
    ]
    for title, text in scenarios:
        doc.add_heading(title, level=2)
        doc.add_paragraph(text)

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════
    # 14. LIMITATIONS & FUTURE ENHANCEMENTS
    # ═══════════════════════════════════════════════════════════════
    doc.add_heading("14. Limitations & Future Enhancements", level=1)

    doc.add_heading("14.1 Current Limitations", level=2)
    add_table(
        doc,
        ["Limitation", "Business Impact"],
        [
            ["No formal approval workflow", "Supervisor sign-off is not system-enforced"],
            ["No edit/delete for logbook entries", "Errors require workarounds or IT support"],
            ["No submitter attribution on logbook rows", "Audits cannot identify who entered a reading"],
            ["No operational alerts/notifications", "Incidents and missed entries are not proactively flagged"],
            ["Manager field unused in workflows", "No escalation or reporting hierarchy in the system"],
            ["Batch submit skips duplicate check", "Potential duplicate daily production records"],
            ["Accident register (ehs_accident) not live", "Form built but not accessible in menu"],
            ["Equipment edits not permission-scoped", "Any logged-in user can modify equipment history"],
            ["No mobile-native app", "Web browser access only"],
            ["BI dashboards lack one-click export", "CSV available from source forms, not dashboards directly"],
        ],
    )

    doc.add_heading("14.2 Requested / Recommended Enhancements", level=2)
    add_bullets(
        doc,
        [
            "Multi-level approval workflow with reviewer queues and status tracking.",
            "Edit/correct committed entries with full audit trail.",
            "Submitter identity captured on every logbook record.",
            "Email/SMS alerts for EHS incidents and critical stoppages.",
            "Missed logbook entry reminders by shift.",
            "Dedicated auditor read-only role with enhanced search.",
            "Mobile-friendly or native mobile application for field entry.",
            "Dashboard PDF/Excel export.",
            "Activate Accident Register (ehs_accident) in the EHS module.",
            "Role-based equipment history edit permissions.",
        ],
    )

    doc.add_heading("14.3 Future Roadmap Opportunities", level=2)
    add_bullets(
        doc,
        [
            "Integration with plant DCS/SCADA for automatic reading import.",
            "Predictive analytics on equipment maintenance history.",
            "Unified executive dashboard across all plant areas.",
            "Digital signature integration for compliance.",
            "Offline capture with sync for connectivity-limited areas.",
        ],
    )

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════
    # 15. GLOSSARY
    # ═══════════════════════════════════════════════════════════════
    doc.add_heading("15. Glossary", level=1)
    glossary = [
        ("BI Control Tower", "Analytics module providing distillery and milling operational dashboards."),
        ("Batch Submit", "Submitting multiple log rows in a single action (production pan/decanter/clarification forms)."),
        ("Confirm & Commit", "Final action in the review modal that permanently saves a logbook entry."),
        ("CSV Export", "Downloading historical form records as a comma-separated file for Excel analysis."),
        ("Dashboard Mapping", "Admin assignment controlling which BI dashboards an employee can access."),
        ("DigiLog", "Zuari Industries digital logbook and operations platform."),
        ("Distillery Operations", "Daily snapshot form capturing ethanol production, efficiencies, and storage."),
        ("DS / RS", "Double Sulphitation / Refinery Sulphitation — sugar manufacturing process stages."),
        ("EHS", "Environment, Health & Safety — incident reporting and water management module."),
        ("Equipment History Card", "Digital record of an asset's specifications, OEM schedule, and maintenance timeline."),
        ("Form Hub", "Central page listing all operational modules assigned to an employee."),
        ("Form Key", "Unique system identifier for each logbook form (e.g., mill_logbook1)."),
        ("Form Mapping", "Admin assignment restricting an employee to specific forms within an app."),
        ("FS%", "Fermentable Sugar percentage — distillery efficiency metric (auto-calculated)."),
        ("GWA", "Ground Water Abstraction — EHS water dashboard for bore well extraction."),
        ("ETP", "Effluent Treatment Plant — EHS water dashboard for treated effluent quality and quantity."),
        ("CPU", "Condensate Polishing Unit — EHS water dashboard for recycled water quality."),
        ("HOD", "Head of Department — referenced in EHS forms as text fields, not a system role."),
        ("Mapping", "Admin configuration linking an employee to an app (and optionally specific forms)."),
        ("MTD", "Month to Date — common date filter preset on BI dashboards."),
        ("Near Miss", "An EHS report for incidents that did not result in injury but had potential."),
        ("OEM Schedule", "Original Equipment Manufacturer recommended maintenance intervals."),
        ("Pol / Brix", "Laboratory measurements of polarization and sugar content."),
        ("Power Plant Hierarchy", "Tree structure organizing plant equipment by boiler, turbine, WTP, etc."),
        ("Review Modal", "Pre-submit summary screen requiring user certification before commit."),
        ("Shift", "Operating shift designation: A, B, C, or G."),
        ("SSO", "Single Sign-On — login via Microsoft or Google corporate accounts."),
        ("Stoppage", "Recorded downtime event with start/end time, section, machinery, and remarks."),
        ("TRS", "Total Recoverable Sugar — key milling/distillery performance metric."),
        ("View Data", "In-app modal for browsing paginated historical submitted records."),
        ("WTP", "Water Treatment Plant — area within the power plant equipment hierarchy."),
    ]
    add_table(doc, ["Term", "Definition"], glossary)

    # Footer note
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("— End of Document —").italic = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUTPUT))
    return OUTPUT


if __name__ == "__main__":
    path = build_document()
    print(f"Generated: {path}")

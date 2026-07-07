"""Generate DigiLog Technical Documentation as Word (.docx) and Markdown (.md)."""
from __future__ import annotations

import sys
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).resolve().parent.parent
SCRIPTS_DIR = Path(__file__).resolve().parent
DOCS_DIR = ROOT / "docs"
OUTPUT_DOCX = DOCS_DIR / "DigiLog_Technical_Documentation.docx"
OUTPUT_MD = DOCS_DIR / "DigiLog_Technical_Documentation.md"
DOC_DATE = date.today().strftime("%d %B %Y")
SCHEMA_SQL = ROOT / "backend" / "backup_before_reconcile.sql"

if str(SCRIPTS_DIR) not in sys.path:
    sys.path.insert(0, str(SCRIPTS_DIR))

from schema_parser import parse_sql_file, column_rows  # noqa: E402


# ─── docx helpers (same pattern as generate_functional_documentation.py) ───────

def set_cell_shading(cell, hex_color: str):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), hex_color)
    cell._tc.get_or_add_tcPr().append(shading)


def add_table(doc, headers, rows, header_fill="1F4E79"):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.size = Pt(10)
        set_cell_shading(hdr[i], header_fill)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            table.rows[ri + 1].cells[ci].text = str(val)
            for p in table.rows[ri + 1].cells[ci].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
    doc.add_paragraph()
    return table


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")
        for run in p.runs:
            run.font.size = Pt(11)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(item, style="List Number")
        for run in p.runs:
            run.font.size = Pt(11)


def add_code_block(doc, text: str):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)


def add_labeled(doc, label: str, text: str):
    p = doc.add_paragraph()
    p.add_run(f"{label}: ").bold = True
    p.add_run(text)


def table_description_map() -> dict[str, str]:
    return {name: desc for name, _domain, desc in DATABASE_TABLES}


def load_table_schemas():
    """Parse live schema from mysqldump snapshot."""
    sql_path = SCHEMA_SQL if SCHEMA_SQL.is_file() else ROOT / "mysql" / "init.sql"
    return parse_sql_file(sql_path)


def render_schema_sections_docx(doc, schemas, table_descriptions: dict[str, str]):
    doc.add_heading("8.4 Complete Table Schemas (Column-Level)", level=2)
    doc.add_paragraph(
        f"Full column definitions for all {len(schemas)} application tables, parsed from "
        f"`{SCHEMA_SQL.relative_to(ROOT)}` (MySQL 8 mysqldump snapshot, July 2026). "
        "Logbook tables typically have no PRIMARY KEY; uniqueness is enforced at application "
        "layer via operational keys (Date + Shift + Time, etc.)."
    )

    current_domain = None
    for schema in schemas:
        if schema.domain != current_domain:
            current_domain = schema.domain
            doc.add_heading(f"Domain: {current_domain}", level=3)

        doc.add_heading(f"Table: `{schema.name}`", level=4)
        desc = table_descriptions.get(schema.name, "")
        if desc:
            add_labeled(doc, "Business purpose", desc)
        meta = f"Engine={schema.engine or 'InnoDB'}"
        if schema.charset:
            meta += f", Charset={schema.charset}"
        add_labeled(doc, "Storage", meta)

        add_table(
            doc,
            ["Column", "Data Type", "Nullable", "Default", "Extra"],
            column_rows(schema),
            header_fill="2E5984",
        )

        constraint_lines = []
        if schema.primary_key:
            constraint_lines.append(f"PRIMARY KEY ({', '.join(schema.primary_key)})")
        constraint_lines.extend(schema.unique_keys)
        constraint_lines.extend(schema.indexes)
        constraint_lines.extend(schema.foreign_keys)
        constraint_lines.extend(schema.constraints)

        if constraint_lines:
            doc.add_paragraph("Keys & constraints:")
            add_bullets(doc, constraint_lines[:20])  # cap bullets for very wide tables
            if len(constraint_lines) > 20:
                doc.add_paragraph(f"... and {len(constraint_lines) - 20} more index/constraint lines.")
        else:
            doc.add_paragraph(
                "Keys: none defined at database level (append-only logbook pattern)."
            )


def render_schema_sections_md(md: "MdBuilder", schemas, table_descriptions: dict[str, str]):
    md.h(3, "8.4 Complete Table Schemas (Column-Level)")
    md.p(
        f"Full column definitions for all **{len(schemas)}** application tables, sourced from "
        f"`backend/backup_before_reconcile.sql`. "
        "Logbook tables typically have no PRIMARY KEY; duplicate prevention uses application-layer "
        "operational keys."
    )

    current_domain = None
    for schema in schemas:
        if schema.domain != current_domain:
            current_domain = schema.domain
            md.h(4, f"Domain: {current_domain}")

        md.h(5, f"Table: `{schema.name}`")
        desc = table_descriptions.get(schema.name, "")
        if desc:
            md.p(f"**Business purpose:** {desc}")
        meta = f"Engine={schema.engine or 'InnoDB'}"
        if schema.charset:
            meta += f", Charset={schema.charset}"
        md.p(f"**Storage:** {meta}")

        md.table(
            ["Column", "Data Type", "Nullable", "Default", "Extra"],
            column_rows(schema),
        )

        constraint_lines = []
        if schema.primary_key:
            constraint_lines.append(f"PRIMARY KEY ({', '.join(schema.primary_key)})")
        constraint_lines.extend(schema.unique_keys)
        constraint_lines.extend(schema.indexes)
        constraint_lines.extend(schema.foreign_keys)
        constraint_lines.extend(schema.constraints)

        if constraint_lines:
            md.p("**Keys & constraints:**")
            md.bullets(constraint_lines)
        else:
            md.p("*Keys: none at database level (append-only logbook pattern).*")


# ─── markdown helpers ────────────────────────────────────────────────────────

class MdBuilder:
    def __init__(self):
        self.lines: list[str] = []

    def h(self, level: int, text: str):
        self.lines.append(f"{'#' * level} {text}\n")

    def p(self, text: str = ""):
        self.lines.append(f"{text}\n")

    def bullets(self, items: list[str]):
        for item in items:
            self.lines.append(f"- {item}")
        self.lines.append("")

    def numbered(self, items: list[str]):
        for i, item in enumerate(items, 1):
            self.lines.append(f"{i}. {item}")
        self.lines.append("")

    def table(self, headers: list[str], rows: list[list[str]]):
        self.lines.append("| " + " | ".join(headers) + " |")
        self.lines.append("| " + " | ".join("---" for _ in headers) + " |")
        for row in rows:
            self.lines.append("| " + " | ".join(str(c) for c in row) + " |")
        self.lines.append("")

    def code(self, text: str, lang: str = ""):
        self.lines.append(f"```{lang}")
        self.lines.append(text)
        self.lines.append("```\n")

    def build(self) -> str:
        return "\n".join(self.lines)


# ─── shared content data ─────────────────────────────────────────────────────

API_ENDPOINTS = [
    # Auth
    ("POST", "/api/auth/login", "Public (login rate-limited)", "Email/password login; adminPortal flag for admin portal"),
    ("POST", "/api/auth/outlook", "Public", "Microsoft SSO — exchanges MSAL access token for JWT"),
    ("POST", "/api/auth/google", "Public", "Google SSO — verifies ID token, returns JWT"),
    ("GET", "/api/auth/me", "Bearer", "Current user profile from JWT"),
    ("GET", "/api/auth/users/:userId/avatar", "Bearer", "Stream user avatar image (authenticated)"),
    ("POST", "/api/auth/me/avatar", "Bearer + multipart", "Upload/crop profile avatar (multer)"),
    ("DELETE", "/api/auth/me/avatar", "Bearer", "Remove profile avatar"),
    # Admin
    ("GET", "/api/admin/users", "Bearer admin", "List all users"),
    ("POST", "/api/admin/users", "Bearer admin", "Create employee/admin user"),
    ("PUT", "/api/admin/users/:id", "Bearer admin", "Update user fields, role, active status"),
    ("DELETE", "/api/admin/users/:id", "Bearer admin", "Delete user"),
    ("PUT", "/api/admin/users/:id/manager", "Bearer admin", "Assign reporting manager (informational)"),
    ("POST", "/api/admin/users/:id/send-mail", "Bearer admin", "Send activation email with temp password (local only)"),
    ("POST", "/api/admin/users/send-mail-bulk", "Bearer admin", "Bulk activation emails"),
    ("GET", "/api/admin/mappings", "Bearer admin", "All user→app mappings with form restrictions"),
    ("POST", "/api/admin/mappings", "Bearer admin", "Upsert mapping + mapping_forms"),
    ("DELETE", "/api/admin/mappings/:id", "Bearer admin", "Remove mapping"),
    ("GET", "/api/admin/apps-all", "Bearer admin", "All apps with nested forms (for mapping UI)"),
    ("GET", "/api/admin/bi-settings", "Bearer admin", "Read portal_settings BI toggles"),
    ("PUT", "/api/admin/bi-settings", "Bearer admin", "Update BI settings (e.g. third season compare)"),
    ("GET", "/api/admin/data-upload-access", "Bearer admin", "List employees with data upload grant"),
    ("PUT", "/api/admin/data-upload-access", "Bearer admin", "Grant/revoke data upload access"),
    # Apps
    ("GET", "/api/apps", "Bearer", "Apps accessible to current user (respects mappings)"),
    # Forms
    ("GET", "/api/forms/:formKey", "Bearer + mapping", "Form metadata (name, description)"),
    ("GET", "/api/forms/:formKey/records", "Bearer + mapping", "Paginated records (?page, ?limit≤10000)"),
    ("POST", "/api/forms/:formKey", "Bearer + mapping", "Single submit with duplicate detection"),
    ("POST", "/api/forms/:formKey/batch", "Bearer + mapping", "Batch insert (prod pan/decanter/clarification)"),
    # Mill House Equipment
    ("GET", "/api/equipment", "Bearer", "List mh_equipment"),
    ("GET", "/api/equipment/:id", "Bearer", "Equipment detail with specs/schedule"),
    ("PUT", "/api/equipment/:id", "Bearer", "Update equipment header fields"),
    ("PUT", "/api/equipment/:id/image/:type", "Bearer + multipart", "Upload photo/plate (type=photo|plate)"),
    ("DELETE", "/api/equipment/:id/image/:type", "Bearer", "Remove equipment image"),
    ("PUT", "/api/equipment/:id/specs", "Bearer", "Replace specification rows"),
    ("PUT", "/api/equipment/:id/schedule", "Bearer", "Replace OEM schedule rows"),
    ("GET", "/api/equipment/:id/history", "Bearer", "Maintenance history timeline"),
    ("POST", "/api/equipment/:id/history", "Bearer", "Add history entry"),
    ("PUT", "/api/equipment/:id/history/:hid", "Bearer", "Update history entry"),
    ("DELETE", "/api/equipment/:id/history/:hid", "Bearer", "Delete history entry"),
    # Legacy Power Plant Equipment
    ("GET", "/api/power/lookup", "Bearer", "Search equipment by tag/name within dept"),
    ("GET", "/api/power", "Bearer", "List pp_equipment (filter by dept query)"),
    ("POST", "/api/power", "Bearer", "Create pp_equipment"),
    ("GET", "/api/power/:id", "Bearer", "Equipment detail"),
    ("PUT", "/api/power/:id", "Bearer", "Update equipment"),
    ("PUT", "/api/power/:id/image/:type", "Bearer + multipart", "Upload image"),
    ("DELETE", "/api/power/:id/image/:type", "Bearer", "Delete image"),
    ("PUT", "/api/power/:id/specs", "Bearer", "Update specs"),
    ("PUT", "/api/power/:id/schedule", "Bearer", "Update OEM schedule"),
    ("GET", "/api/power/:id/history", "Bearer", "History list"),
    ("POST", "/api/power/:id/history", "Bearer", "Add history"),
    ("PUT", "/api/power/:id/history/:hid", "Bearer", "Update history"),
    ("DELETE", "/api/power/:id/history/:hid", "Bearer", "Delete history"),
    # Power Plant Equipment New (PPN)
    ("GET", "/api/power-new/lookup", "Bearer", "Hierarchy-aware equipment lookup"),
    ("GET", "/api/power-new/hierarchy", "Bearer", "Full ppn_hierarchy_node tree"),
    ("GET", "/api/power-new/hierarchy/path/:nodeId", "Bearer", "Breadcrumb path to node"),
    ("POST", "/api/power-new/hierarchy", "Bearer", "Create hierarchy node"),
    ("PUT", "/api/power-new/hierarchy/:nodeId", "Bearer", "Rename/reparent node (protected seeds blocked)"),
    ("DELETE", "/api/power-new/hierarchy/:nodeId", "Bearer", "Delete node (protected seeds blocked)"),
    ("GET", "/api/power-new", "Bearer", "List ppn_equipment"),
    ("POST", "/api/power-new", "Bearer", "Create ppn_equipment"),
    ("GET", "/api/power-new/:id", "Bearer", "Equipment detail with section-scoped specs"),
    ("PUT", "/api/power-new/:id", "Bearer", "Update equipment"),
    ("PUT", "/api/power-new/:id/image/:type", "Bearer + multipart", "Upload image"),
    ("DELETE", "/api/power-new/:id/image/:type", "Bearer", "Delete image"),
    ("PUT", "/api/power-new/:id/specs", "Bearer", "Update specs (section/sub_section)"),
    ("PUT", "/api/power-new/:id/schedule", "Bearer", "Update OEM schedule"),
    ("DELETE", "/api/power-new/:id/history-sub-group", "Bearer", "Delete scoped history sub-group"),
    ("PUT", "/api/power-new/:id/history-sub-group/rename", "Bearer", "Rename history sub-group"),
    ("GET", "/api/power-new/:id/history", "Bearer", "History (optionally scoped by section)"),
    ("POST", "/api/power-new/:id/history", "Bearer", "Add history entry"),
    ("PUT", "/api/power-new/:id/history/:hid", "Bearer", "Update history"),
    ("DELETE", "/api/power-new/:id/history/:hid", "Bearer", "Delete history"),
    # BI
    ("GET", "/api/bi/settings", "Bearer", "Employee-readable BI portal settings"),
    ("GET", "/api/bi/distillery-operations", "Bearer + mapping", "Distillery analytics (?from, ?to ISO dates)"),
    ("GET", "/api/bi/milling-operations", "Bearer + mapping", "Mill stoppage/outage analytics"),
    ("GET", "/api/bi/milling-equipment-temp", "Bearer + mapping", "Equipment temperature BI series"),
    ("GET", "/api/bi/milling-shredder", "Bearer + mapping", "Shredder/OTG BI series"),
    ("GET", "/api/bi/milling-lube-roller", "Bearer + mapping", "Lube pressure & roller temp BI series"),
    # Homepage
    ("GET", "/api/homepage-cards", "Bearer", "user_homepage_cards for current user"),
    # Data Upload
    ("GET", "/api/data-upload/access", "Bearer", "Whether current user has upload access"),
    ("GET", "/api/data-upload/files", "Bearer + upload access", "List uploaded files with uploader audit"),
    ("POST", "/api/data-upload", "Bearer + upload access + multipart", "Upload CSV/XLS/XLSX (max DATA_UPLOAD_MAX_BYTES)"),
    ("GET", "/api/data-upload/files/:id/download", "Bearer + upload access", "Download stored file"),
    ("DELETE", "/api/data-upload/files/:id", "Bearer + upload access", "Delete own upload"),
    # Health
    ("GET", "/api/health", "Public", "Liveness probe { status: ok }"),
]

FORM_CONFIG_ROWS = [
    ("mill_logbook1", "mill_logbook1", "A", "Date+Shift+Time", "Equipment Temperature"),
    ("mill_logbook2", "mill_logbook2", "A", "Date+Shift+Time", "Shredder / OTG"),
    ("mill_logbook3", "mill_logbook3", "A", "Date+Shift+Time", "Lube Pressure / Roller Temp"),
    ("mill_stoppages", "mill_stoppages", "B", "Date+start+end", "Mill Stoppages"),
    ("ds_logbook", "ds_logbook", "C", "Date+Shift+Sampling_time", "DS Logbook"),
    ("rs_logbook", "rs_logbook", "C", "Date+Shift+Sampling_time", "RS Logbook"),
    ("ops_logbook", "ops_logbook", "C", "Date+Shift+Sampling_time", "Operations Logbook"),
    ("sa_logbook", "sa_logbook", "C", "Date+Shift+Sampling_time", "Special Analysis (tsCol=timestamp_col)"),
    ("syrp_logbook", "syrp_logbook", "D", "Date+Shift", "Syrup Logbook"),
    ("stoppage_logbook", "stoppage_logbook", "B", "Date+start+end", "Lab Stoppage"),
    ("ph_power", "ph_power", "E", "Date+Time", "Power Details"),
    ("ph_steam", "ph_steam", "E", "Date+Time", "Steam Details"),
    ("ph_stoppage", "ph_stoppage", "B", "Date+start+end", "Power Stoppage"),
    ("distillery_ops", "distillery_operations", "G", "Date only", "Distillery daily snapshot; excludes generated FS%, total_mol_in_store_qtls"),
    ("ehs_near_miss", "ehs_near_miss", "E", "Date+Time", "Near Miss / Incident"),
    ("ehs_accident", "ehs_accident", "E", "Date+Time", "Accident register (table exists; UI route not wired)"),
    ("ehs_water_gwa", "ehs_water_gwa", "G", "Date only", "Ground Water Abstraction"),
    ("ehs_water_etp", "ehs_water_etp", "G", "Date only", "ETP Working; pH 0–14 validation"),
    ("ehs_water_cpu", "ehs_water_cpu", "G", "Date only", "CPU Water Recycle; pH 0–14 validation"),
    ("prod_shift_chemist", "prod_shift_chemist", "G", "Date only", "Shift Chemist Job Log"),
    ("prod_centrifugal", "prod_centrifugal", "H", "Date+Shift", "A-Centrifugal Stoppage"),
    ("prod_pan_logbook", "prod_pan_logbook", "G", "Date only", "Pan Log Book — batch via /batch"),
    ("prod_decanter", "prod_decanter", "G", "Date only", "Decanter — batch via /batch (24 slots)"),
    ("prod_clarification", "prod_clarification", "G", "Date only", "Clarification — batch via /batch"),
]

PATTERN_DESCRIPTIONS = [
    ("A", "Date, Shift, Time", "Mill logbooks (equipment temp, shredder, lube)"),
    ("B", "Date, start_time, end_time", "Stoppage forms (mill, lab, power)"),
    ("C", "Date, Shift, Sampling_time", "Lab logbooks (DS, RS, Ops, SA)"),
    ("D", "Date, Shift", "Syrup logbook (no time column)"),
    ("E", "Date, Time", "Power logbooks; EHS near miss/accident"),
    ("G", "Date only", "Daily snapshot (distillery, EHS water, production daily)"),
    ("H", "Date, Shift", "Production centrifugal stoppage"),
]

DATABASE_TABLES = [
    # System
    ("users", "System", "Accounts: email, bcrypt password, role admin|employee, auth_provider local|outlook|google, manager_id, avatar"),
    ("apps", "System", "Registered applications (Mill Logbook, Lab, BI Control Tower, etc.)"),
    ("forms", "System", "Form registry linked to apps via app_id; form_key drives FORM_CONFIG"),
    ("mappings", "System", "user_id + app_id — grants app access to employee"),
    ("mapping_forms", "System", "Optional per-form restriction within a mapping; empty = all forms in app"),
    ("portal_settings", "System", "Key/value admin settings (bi_third_season_compare)"),
    ("user_homepage_cards", "System", "card_key forms_hub | bi_control_tower for /dashboard"),
    ("user_data_upload_access", "System", "Admin-granted flag for Data Ingestion Center"),
    ("data_upload_files", "System", "Uploaded file registry with uploader audit trail"),
    # Mill
    ("mill_logbook1", "Mill", "Equipment motor/gear/bearing temperatures by mill section"),
    ("mill_logbook2", "Mill", "Shredder vibration/OTG readings"),
    ("mill_logbook3", "Mill", "Lube pressures and roller temperatures"),
    ("mill_stoppages", "Mill", "Mill downtime events: section, machinery, remarks"),
    # Lab
    ("ds_logbook", "Lab", "Double Sulphitation pol/brix measurements"),
    ("rs_logbook", "Lab", "Refinery Sulphitation analysis incl. IU/pH fields"),
    ("ops_logbook", "Lab", "Operations crush, imbibe, bagging, FBD readings"),
    ("sa_logbook", "Lab", "Special analysis retention/moisture/colour"),
    ("syrp_logbook", "Lab", "Syrup production and TRS metrics"),
    ("stoppage_logbook", "Lab", "Lab department stoppages"),
    # Power logbooks
    ("ph_power", "Power", "Generation, export, import, consumption by unit"),
    ("ph_steam", "Power", "Steam generation and consumption balance"),
    ("ph_stoppage", "Power", "Power house stoppages with category"),
    # Legacy power equipment
    ("pp_equipment", "Power Equipment", "Legacy dept-based equipment cards (electrical/instrument/mechanical)"),
    ("pp_specs", "Power Equipment", "Label/value specification rows per pp_equipment"),
    ("pp_oem_schedule", "Power Equipment", "OEM maintenance interval matrix"),
    ("pp_history", "Power Equipment", "Maintenance timeline with before/after images"),
    # New power equipment (PPN)
    ("ppn_equipment", "Power Equipment New", "Hierarchy-linked equipment records"),
    ("ppn_specs", "Power Equipment New", "Section/sub_section scoped specifications"),
    ("ppn_oem_schedule", "Power Equipment New", "Scoped OEM schedule with equipment_refs JSON"),
    ("ppn_history", "Power Equipment New", "Scoped maintenance history with equipment_refs"),
    ("ppn_hierarchy_node", "Power Equipment New", "Tree nodes (group|equipment) for 150TPH/70TPH/WTP areas"),
    # Mill house equipment
    ("mh_equipment", "Mill House Equipment", "Mill house asset registry"),
    ("mh_specs", "Mill House Equipment", "Equipment specifications"),
    ("mh_oem_schedule", "Mill House Equipment", "OEM schedule per asset"),
    ("mh_history", "Mill House Equipment", "Maintenance history with images"),
    # Distillery
    ("distillery_operations", "Distillery", "Daily ops snapshot; generated columns FS%, total_mol_in_store_qtls"),
    # EHS
    ("ehs_near_miss", "EHS", "Near miss/incident reports with HOD text fields"),
    ("ehs_accident", "EHS", "Accident register (not exposed in App.jsx menu)"),
    ("ehs_water_gwa", "EHS", "Ground water abstraction meters and allocation"),
    ("ehs_water_etp", "EHS", "ETP inlet/outlet quality and flow"),
    ("ehs_water_cpu", "EHS", "CPU recycle water quality (pH per shift)"),
    # Production
    ("prod_shift_chemist", "Production", "Shift chemist instructions and job logs"),
    ("prod_centrifugal", "Production", "Centrifugal machine stoppages M1–M4"),
    ("prod_pan_logbook", "Production", "Pan strike records (batch submit)"),
    ("prod_decanter", "Production", "Decanter hourly readings ST1/ST2 (batch submit)"),
    ("prod_clarification", "Production", "Clarification process readings (batch submit)"),
    # BI reference data
    ("data_mill_mapping", "BI Reference", "Mill thermal variable→equipment mapping for BI reports"),
    ("data_shredder_mapping", "BI Reference", "Shredder variable mapping"),
    ("data_lube_mapping", "BI Reference", "Lube/roller variable mapping"),
]

ARCH_ASCII = """\
┌─────────────────────────────────────────────────────────────────────────┐
│                         CLIENT TIER (Browser)                           │
│  React 18 + Vite 6 + Tailwind CSS + React Router 6                      │
│  Axios (JWT interceptor) │ MSAL (@azure/msal-react) │ Google OAuth      │
│  Routes: /dashboard, /forms-hub, /bi/*, /forms/*, /equipment, /power*   │
└───────────────────────────────┬─────────────────────────────────────────┘
                                │ HTTPS  JSON  Authorization: Bearer <JWT>
                                │ Base URL: VITE_API_URL or same-origin /api
                                ▼
┌─────────────────────────────────────────────────────────────────────────┐
│                    API TIER — Express 4.19 (Node.js)                    │
│  helmet │ cors(CLIENT_ORIGIN) │ express-rate-limit │ json 10mb          │
│  /api/auth │ /api/admin │ /api/forms │ /api/equipment │ /api/power*     │
│  /api/bi │ /api/homepage-cards │ /api/data-upload │ /api/health         │
│  middleware: authenticate → requireRole(admin) → canAccessForm          │
└───────────────────────────────┬─────────────────────────────────────────┘
                                │ mysql2/promise connection pool
                                │ DB_POOL_LIMIT=30, queueLimit=100
                                ▼
┌─────────────────────────────────────────────────────────────────────────┐
│                         DATA TIER — MySQL 8                             │
│  Runtime: raw SQL via mysql2 (NOT Prisma Client)                        │
│  Schema: mysql/init.sql + migrate_*.sql + Prisma migrations (history)  │
│  ~49 tables: users, forms, logbooks, equipment, BI mappings             │
└─────────────────────────────────────────────────────────────────────────┘"""

ARCH_MERMAID = """\
graph TB
    subgraph Client["Frontend — React 18 / Vite 6"]
        SPA[SPA Routes App.jsx]
        AuthCtx[AuthContext + MSAL + Google]
        Forms[Form Pages + FormReviewModal]
        BI[BI Dashboards Recharts]
        SPA --> AuthCtx
        SPA --> Forms
        SPA --> BI
    end

    subgraph API["Backend — Express 4.19"]
        AuthR[/api/auth JWT Bearer/]
        AdminR[/api/admin role=admin/]
        FormR[/api/forms FORM_CONFIG/]
        EquipR[/api/equipment + /api/power*/]
        BiR[/api/bi date bounds/]
        AuthR --> MW[authenticate middleware]
        AdminR --> MW
        FormR --> MW
        EquipR --> MW
        BiR --> MW
    end

    subgraph Data["MySQL 8"]
        Pool[(mysql2 pool)]
        Tables[(init.sql tables)]
        Pool --> Tables
    end

    AuthCtx -->|Axios Bearer| AuthR
    Forms -->|POST submit| FormR
    BI -->|GET ?from&to| BiR
    MW --> Pool"""


def build_docx() -> Path:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Cover
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("DigiLog")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(31, 78, 121)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = sub.add_run("Technical Documentation")
    r2.bold = True
    r2.font.size = Pt(18)
    r2.font.color.rgb = RGBColor(68, 84, 106)

    doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run("Zuari Industries — Digital Operations Platform\n").font.size = Pt(12)
    meta.add_run("Document Version: 1.0\n").font.size = Pt(11)
    meta.add_run(f"Date: {DOC_DATE}\n").font.size = Pt(11)
    meta.add_run(
        "Audience: Developers, DevOps Engineers, Security Auditors, System Maintainers"
    ).font.size = Pt(11)
    doc.add_page_break()

    # Document control
    doc.add_heading("Document Control", level=1)
    add_table(
        doc,
        ["Item", "Details"],
        [
            ["Document Title", "DigiLog Technical Documentation"],
            ["System Name", "DigiLog — Digital Logbook Platform"],
            ["Organization", "Zuari Industries"],
            ["Purpose", "Technical reference for architecture, APIs, database, deployment, and maintenance"],
            ["Classification", "Internal — Technical Use"],
            ["Repository", str(ROOT)],
        ],
    )
    doc.add_page_break()

    # ── 1. Project Overview ──
    doc.add_heading("1. Project Overview", level=1)
    doc.add_paragraph(
        "DigiLog is Zuari Industries' plant-wide digital logbook and operations platform. "
        "It replaces paper registers with a React single-page application backed by an Express REST API "
        "and MySQL 8 database. Operators submit shift readings across mill, laboratory, power house, "
        "distillery, EHS, and production areas. Engineering teams maintain digital equipment history cards. "
        "Managers consume BI dashboards built from the same operational data."
    )
    doc.add_heading("1.1 Repository Layout", level=2)
    add_bullets(
        doc,
        [
            f"{ROOT.name}/frontend — React 18 SPA (Vite dev server port 5173)",
            f"{ROOT.name}/backend — Express API (default port 5000)",
            f"{ROOT.name}/mysql — init.sql baseline schema and migrate_*.sql patches",
            f"{ROOT.name}/scripts — Python/Node utilities including this documentation generator",
            f"{ROOT.name}/docs — Generated documentation outputs",
        ],
    )
    doc.add_heading("1.2 Key Technical Characteristics", level=2)
    add_table(
        doc,
        ["Characteristic", "Implementation"],
        [
            ["API style", "REST JSON under /api/* prefix"],
            ["Auth", "JWT Bearer tokens (7d default expiry); roles admin | employee"],
            ["Data access", "mysql2 raw SQL at runtime; Prisma used for migrations only"],
            ["Access control", "mappings + mapping_forms + user_homepage_cards + user_data_upload_access"],
            ["Form workflow", "FormReviewModal = pre-submit self-certification only (no approval queue)"],
            ["Deployment", "No Docker in repo; Node + MySQL + static Vite build"],
        ],
    )
    doc.add_page_break()

    # ── 2. System Architecture ──
    doc.add_heading("2. System Architecture", level=1)
    doc.add_heading("2.1 Three-Tier Architecture", level=2)
    add_code_block(doc, ARCH_ASCII)
    doc.add_heading("2.2 Request Flow — Form Submit", level=2)
    add_numbered(
        doc,
        [
            "Employee opens /forms/<formKey> — React page loads form fields.",
            "User clicks Submit → FormReviewModal shows read-only summary.",
            "Confirm & Commit → POST /api/forms/:formKey with JSON body.",
            "Axios interceptor attaches Authorization: Bearer <JWT>.",
            "authenticate middleware validates token and loads users row.",
            "form.controller canAccessForm checks mappings/mapping_forms.",
            "injectDateCols applies pattern A–H date/shift/time columns.",
            "hasDuplicateOperationRow blocks duplicate operational keys (single submit).",
            "INSERT into target table; 201 response; toast on success.",
        ],
    )
    doc.add_heading("2.3 Portal Separation", level=2)
    doc.add_paragraph(
        "auth.controller enforceAdminPortalRules blocks cross-portal login: admins cannot use the home "
        "portal (/), employees cannot use /admin/login. SSO users must be pre-provisioned in users table."
    )
    doc.add_page_break()

    # ── 3. Technology Stack ──
    doc.add_heading("3. Technology Stack", level=1)
    add_table(
        doc,
        ["Layer", "Technology", "Version / Notes"],
        [
            ["Frontend framework", "React", "18.3.x"],
            ["Build tool", "Vite", "6.4.x"],
            ["Styling", "Tailwind CSS", "3.4.x + PostCSS + Autoprefixer"],
            ["Routing", "React Router DOM", "6.30.x"],
            ["HTTP client", "Axios", "1.7.x with JWT interceptor"],
            ["Microsoft SSO", "@azure/msal-browser + msal-react", "3.14 / 2.0"],
            ["Google SSO", "@react-oauth/google", "0.12.x"],
            ["Charts", "Recharts", "3.8.x (BI dashboards)"],
            ["Notifications", "react-hot-toast", "2.4.x"],
            ["Backend runtime", "Node.js + Express", "4.19.x"],
            ["Database driver", "mysql2", "3.10.x promise pool"],
            ["ORM (migrations only)", "Prisma", "7.8.x — models @@ignore, runtime uses mysql2"],
            ["Auth", "jsonwebtoken + bcryptjs", "JWT HS256; bcrypt password hash"],
            ["Security", "helmet + express-rate-limit + cors", "Global 500/15min; login 30/15min"],
            ["File upload", "multer", "Avatars + data ingestion files"],
            ["Email", "nodemailer", "Account activation emails (SMTP env vars)"],
            ["Google token verify", "google-auth-library", "Backend Google SSO"],
            ["Spreadsheet", "@e965/xlsx", "Equipment/BI import scripts"],
            ["Database", "MySQL", "8.x utf8mb4_0900_ai_ci"],
        ],
    )
    doc.add_page_break()

    # ── 4. Folder Structure ──
    doc.add_heading("4. Folder Structure", level=1)
    add_table(
        doc,
        ["Path", "Purpose"],
        [
            ["frontend/src/App.jsx", "All protected/public React routes"],
            ["frontend/src/pages/", "Page components (forms, BI, admin, equipment, power)"],
            ["frontend/src/components/", "Shared UI: Navbar, FormReviewModal, ProtectedRoute, BI charts"],
            ["frontend/src/hooks/", "useAuth, useFormMeta, usePowerPlantHierarchy, useDataUploadAccess"],
            ["frontend/src/context/AuthContext.jsx", "JWT session, MSAL/Google login orchestration"],
            ["frontend/src/api/axios.js", "API base URL resolution + Bearer interceptor"],
            ["frontend/src/config/", "formColumnSchemas, powerPlantEquipmentHierarchy"],
            ["backend/server.js", "Express app bootstrap, middleware, route mounting"],
            ["backend/routes/", "10 route modules mounting /api/* endpoints"],
            ["backend/controllers/", "Business logic: auth, form, bi, admin, equipment, power"],
            ["backend/middleware/", "auth.js, roleCheck.js, avatarUpload, dataUploadMulter"],
            ["backend/config/", "env.js, mysql.js, databaseName.js"],
            ["backend/utils/", "jwt, httpError, userPublic, avatarFile, ppnHierarchyLib"],
            ["backend/prisma/", "schema.prisma + migrations/ (schema history)"],
            ["backend/scripts/", "DB seed, import, hierarchy seed, backlog CSV import"],
            ["mysql/init.sql", "Baseline CREATE TABLE for all core tables"],
            ["mysql/migrate_*.sql", "23 incremental SQL patches (EHS, PPN, production, BI)"],
        ],
    )
    doc.add_page_break()

    # ── 5. Frontend Documentation ──
    doc.add_heading("5. Frontend Documentation", level=1)
    doc.add_heading("5.1 Route Map (App.jsx)", level=2)
    add_table(
        doc,
        ["Route", "Component", "Access"],
        [
            ["/", "MarketingLanding (or redirect /dashboard)", "Public"],
            ["/admin/login", "AdminLogin", "Public"],
            ["/dashboard", "HomeLanding", "Protected — homepage cards"],
            ["/forms-hub", "Dashboard", "Protected — assigned apps"],
            ["/bi", "BiControlTower", "Protected"],
            ["/bi/distillery-operations", "DistilleryAnalyticsDashboard", "Protected + BI mapping"],
            ["/bi/milling-operations", "MillingOperationsDashboard", "Protected + BI mapping"],
            ["/apps/:appId", "AppDetail", "Protected"],
            ["/forms/mill_logbook1", "EquipmentTemp", "Protected + form mapping"],
            ["/forms/mill_logbook2", "ShreddarOTG", "Protected + form mapping"],
            ["/forms/mill_logbook3", "LubePressure", "Protected + form mapping"],
            ["/forms/mill_stoppages", "MillStoppages", "Protected + form mapping"],
            ["/forms/ds_logbook … syrp_logbook", "Lab form pages (6)", "Protected + form mapping"],
            ["/forms/ph_power … ph_stoppage", "Power form pages (3)", "Protected + form mapping"],
            ["/forms/distillery_ops", "DistilleryOperations", "Protected + form mapping"],
            ["/equipment, /equipment/:id", "EquipmentList, EquipmentDetail", "Protected"],
            ["/power, /power/:dept/:id", "PowerLanding, PowerList, PowerEquipmentDetail", "Protected"],
            ["/power-plant-equipment-new", "PowerPlantEquipmentNew", "Protected"],
            ["/ehs, /forms/ehs_*", "EhsLanding + 4 live EHS forms", "Protected + form mapping"],
            ["/production, /forms/prod_*", "ProductionLanding + 5 forms", "Protected + form mapping"],
            ["/data-upload", "DataIngestionCenter", "Protected + upload access"],
            ["/admin/employees", "EmployeeManagement", "Protected admin only"],
            ["/operations-desk", "MarketingDashboard", "Public marketing"],
        ],
    )
    doc.add_heading("5.2 Components", level=2)
    add_bullets(
        doc,
        [
            "ProtectedRoute — wraps routes; optional requiredRole='admin'; redirects unauthenticated users",
            "FormReviewModal — pre-submit review/certification; NOT an approval workflow",
            "FormTable + FormPageHeader — shared form layout and View Data / CSV export",
            "Navbar + MobileNavDrawer — authenticated navigation",
            "DigiLogLoginModal + GoogleSignInButton — home portal login (email, MS, Google)",
            "AuthenticatedImage — fetches avatar via Bearer-authenticated blob URL",
            "Equipment* hubs — SpecificationHub, OemMaintenanceScheduleHub, MaintenanceHistoryHub",
            "PowerPlantHierarchyExplorer + HierarchyManagePanel — PPN tree navigation/management",
            "Bi* components — DistilleryChartsGrid, MillRawDataTable, chart expand modals",
            "admin/* modals — EmployeeFormMappingModal, BiDashboardSettings, DataUploadAccessModal",
        ],
    )
    doc.add_heading("5.3 Hooks & Context", level=2)
    add_table(
        doc,
        ["Module", "Purpose"],
        [
            ["context/AuthContext.jsx", "JWT storage, loginManual, loginGoogle, MSAL redirect, refreshUser, logout"],
            ["hooks/useAuth.js", "Consumes AuthContext"],
            ["hooks/useFormMeta.js", "GET /api/forms/:formKey metadata"],
            ["hooks/useAppName.js", "Resolve app display name from apps list"],
            ["hooks/useGsmaFormReview.js", "Form submit → review modal → API POST orchestration"],
            ["hooks/usePowerPlantHierarchy.js", "Fetch/cache PPN hierarchy tree"],
            ["hooks/useDataUploadAccess.js", "GET /api/data-upload/access gate for Data Ingestion Center"],
            ["hooks/useOpenLoginFromQuery.js", "Open login modal when ?login=1 in URL"],
        ],
    )
    doc.add_page_break()

    # ── 6. Backend Documentation ──
    doc.add_heading("6. Backend Documentation", level=1)
    doc.add_heading("6.1 Server Bootstrap (server.js)", level=2)
    add_bullets(
        doc,
        [
            "Loads env.js (JWT_SECRET min 32 chars — process exits if invalid)",
            "testMysqlConnection on startup",
            "helmet (CSP off for JSON API), cors(CLIENT_ORIGIN), json/urlencoded 10mb",
            "globalLimiter 500 req/15min; loginLimiter 30 req/15min on /api/auth/login",
            "Mounts 10 route prefixes under /api/*",
            "Global error handler: mapDbError + logServerError",
        ],
    )
    doc.add_heading("6.2 Controllers", level=2)
    add_table(
        doc,
        ["Controller", "Responsibility"],
        [
            ["auth.controller.js", "login, outlookLogin, googleLogin, getMe, avatar CRUD"],
            ["admin.controller.js", "User CRUD, mappings, activation emails, apps-all"],
            ["form.controller.js", "FORM_CONFIG, submitForm, submitBatch, getRecords, canAccessForm"],
            ["app.controller.js", "getAccessibleApps filtered by mappings"],
            ["equipment.controller.js", "Mill house mh_* CRUD"],
            ["power.controller.js", "Legacy pp_* equipment factory"],
            ["powerNew.controller.js", "PPN equipment with section-scoped history"],
            ["ppnHierarchy.controller.js", "ppn_hierarchy_node tree CRUD with seed protection"],
            ["bi.controller.js", "Distillery/milling BI aggregations with date bounds"],
            ["biSettings.controller.js", "portal_settings BI toggles"],
            ["homepageCards.controller.js", "user_homepage_cards for dashboard"],
            ["dataUpload.controller.js", "File upload registry, access control, mill mapping auto-sync"],
        ],
    )
    doc.add_heading("6.3 Middleware Chain", level=2)
    add_numbered(
        doc,
        [
            "authenticate — Bearer JWT → verifyToken → users lookup → req.user",
            "requireRole('admin') — admin.routes.js router.use guard",
            "canAccessForm — form routes check mappings/mapping_forms",
            "requireDataUploadAccess — data upload routes after getMyAccess",
            "uploadAvatarMiddleware / uploadDataFileMiddleware — multer disk storage",
        ],
    )
    doc.add_page_break()

    # ── 7. API Documentation ──
    doc.add_heading("7. API Documentation", level=1)
    doc.add_paragraph(
        "All API routes are prefixed /api. Authentication uses Authorization: Bearer <JWT> unless noted Public. "
        "Admin routes require role=admin. Form routes additionally enforce mapping access via canAccessForm."
    )
    add_table(
        doc,
        ["Method", "Endpoint", "Auth", "Description"],
        API_ENDPOINTS,
    )
    doc.add_page_break()

    # ── 8. Database Documentation ──
    doc.add_heading("8. Database Documentation", level=1)
    doc.add_heading("8.1 Schema Management", level=2)
    add_bullets(
        doc,
        [
            "Runtime queries use mysql2 pool — NOT Prisma Client (Prisma models are @@ignore)",
            "Baseline: mysql/init.sql applied via npm run db:schema (apply-init-sql.js substitutes DB name)",
            "Incremental: 23 migrate_*.sql files for EHS, production, PPN hierarchy, BI mappings",
            "Prisma migrations in backend/prisma/migrations/ track form/logbook evolution",
            "Schema drift risk: init.sql, migrate_*.sql, and Prisma history can diverge — verify before deploy",
            "Reconciliation: backup_before_reconcile.sql in backend/ is a point-in-time schema snapshot",
        ],
    )
    doc.add_heading("8.2 Tables (49 total)", level=2)
    add_table(doc, ["Table", "Domain", "Description"], DATABASE_TABLES)
    doc.add_heading("8.3 Access Control Tables", level=2)
    add_table(
        doc,
        ["Table", "Relationship"],
        [
            ["mappings", "user_id → app_id (one row per user per app)"],
            ["mapping_forms", "Optional form_id restriction; empty mapping_forms = all forms in app"],
            ["user_homepage_cards", "Controls /dashboard destination cards (forms_hub, bi_control_tower)"],
            ["user_data_upload_access", "Grants /data-upload tab; admins always have access"],
        ],
    )
    schemas = load_table_schemas()
    render_schema_sections_docx(doc, schemas, table_description_map())
    doc.add_page_break()

    # ── 9. Authentication & Authorization ──
    doc.add_heading("9. Authentication & Authorization", level=1)
    doc.add_heading("9.1 JWT Flow", level=2)
    add_numbered(
        doc,
        [
            "Successful login/SSO returns { token, user } — token stored in localStorage",
            "signToken({ id, role }) — JWT_SECRET from env, default expiry JWT_EXPIRES_IN=7d",
            "Every API call: Axios attaches Bearer token",
            "401 on expired/invalid token → localStorage cleared → redirect /?login=1",
        ],
    )
    doc.add_heading("9.2 Roles & Portals", level=2)
    add_table(
        doc,
        ["Role", "Portal", "Capabilities"],
        [
            ["employee", "Home (/)", "Assigned forms, BI dashboards, equipment (any authenticated user)"],
            ["admin", "/admin/login", "Employee management only — cannot access employee portal"],
        ],
    )
    doc.add_heading("9.3 SSO Pre-Provisioning", level=2)
    doc.add_paragraph(
        "Microsoft (outlook) and Google logins require an existing users row matched by email. "
        "First SSO login links microsoft_id/google_id and sets auth_provider. "
        "Unknown emails receive 403: 'You do not have access to use this application.'"
    )
    doc.add_heading("9.4 Form Access Guard", level=2)
    doc.add_paragraph(
        "canAccessForm (form.controller.js): admins bypass; employees need mappings row for the form's app_id; "
        "if mapping_forms has rows, form_id must be listed. Equipment routes have no mapping guard."
    )
    doc.add_page_break()

    # ── 10. Business Logic ──
    doc.add_heading("10. Business Logic", level=1)
    doc.add_heading("10.1 FORM_CONFIG Patterns A–H", level=2)
    add_table(doc, ["Pattern", "Key Columns", "Used By"], PATTERN_DESCRIPTIONS)
    add_table(
        doc,
        ["formKey", "Table", "Pattern", "Duplicate Key", "Notes"],
        FORM_CONFIG_ROWS,
    )
    doc.add_heading("10.2 Duplicate Detection", level=2)
    doc.add_paragraph(
        "Single submit (POST /api/forms/:formKey) calls hasDuplicateOperationRow before INSERT. "
        "Uses NULL-safe <=> comparison on pattern-specific columns. Returns 409 with DUPLICATE_OPERATION_MSG. "
        "Batch submit (POST /api/forms/:formKey/batch) skips duplicate check — used by prod_pan_logbook, "
        "prod_decanter, prod_clarification for multiple rows per session."
    )
    doc.add_heading("10.3 FormReviewModal", level=2)
    doc.add_paragraph(
        "Frontend-only pre-submit review. User certifies accuracy then commits. "
        "No pending/approved/rejected states, no reviewer queues, no manager sign-off gates."
    )
    doc.add_heading("10.4 BI Date Bounds", level=2)
    doc.add_paragraph(
        "bi.controller buildDateBound: default 365-day lookback (DEFAULT_LOOKBACK_DAYS); "
        "accepts ?from=&to= ISO dates; BI_ROW_LIMIT=200000 safety cap on query results."
    )
    doc.add_heading("10.5 Data Upload Auto-Sync", level=2)
    doc.add_paragraph(
        "Recognized mill mapping filenames trigger auto-import into data_mill_mapping / "
        "data_shredder_mapping / data_lube_mapping tables on upload."
    )
    doc.add_page_break()

    # ── 11. Logging & Audit ──
    doc.add_heading("11. Logging & Audit", level=1)
    add_table(
        doc,
        ["Mechanism", "What Is Logged", "Limitation"],
        [
            ["logServerError (httpError.js)", "Stack traces to stderr on server errors", "No centralized log aggregation in repo"],
            ["Logbook timestamp columns", "System timestamp on INSERT (per-table tsCol)", "No submitted_by user_id on logbook rows"],
            ["data_upload_files", "user_id, category, filename, size, created_at", "Full uploader audit for uploads"],
            ["mh_history / pp_history / ppn_history", "created_at, updated_at on maintenance entries", "No edit history trail beyond updated_at"],
            ["users table", "created_at, updated_at, mail_sent flag", "Manager_id is informational only"],
            ["Console startup", "MySQL connection status, server port", "Production should capture stdout/stderr"],
        ],
    )
    doc.add_page_break()

    # ── 12. Deployment & Environment ──
    doc.add_heading("12. Deployment & Environment Setup", level=1)
    doc.add_paragraph(
        "No Docker or docker-compose in the repository. Typical deployment: MySQL 8 instance, "
        "Node backend (PORT), Vite static build served by nginx with /api reverse proxy to Node :5000."
    )
    doc.add_heading("12.1 Backend Environment Variables (env.js)", level=2)
    add_table(
        doc,
        ["Variable", "Default", "Purpose"],
        [
            ["PORT", "5000", "Express listen port"],
            ["NODE_ENV", "development", "production enables trust proxy for rate-limit/HSTS"],
            ["MYSQL_HOST / PORT / USER / PASSWORD", "localhost:3306 root", "MySQL connection (or use DATABASE_URL)"],
            ["MYSQL_DATABASE / DATABASE_URL", "resolved by databaseName.js", "Database name normalization"],
            ["JWT_SECRET", "(required, min 32 chars)", "HS256 signing key — FATAL exit if missing"],
            ["JWT_EXPIRES_IN", "7d", "Token lifetime"],
            ["SMTP_HOST/PORT/USER/PASS/FROM", "empty", "Nodemailer activation emails"],
            ["GOOGLE_CLIENT_ID", "empty", "Google ID token verification"],
            ["CLIENT_ORIGIN", "http://localhost:5173", "CORS allowed origin"],
            ["APP_LOGO_URL", "empty", "Email template logo (falls back to CLIENT_ORIGIN/logo.png)"],
            ["DATA_UPLOAD_MAX_BYTES", "26214400 (25 MB)", "Max upload file size"],
            ["DB_POOL_LIMIT", "30", "mysql2 connectionLimit"],
            ["DB_POOL_QUEUE_LIMIT", "100", "Bounded connection queue"],
            ["DB_CONNECT_TIMEOUT", "10000", "Connection timeout ms"],
        ],
    )
    doc.add_heading("12.2 Frontend Environment Variables", level=2)
    add_table(
        doc,
        ["Variable", "Purpose"],
        [
            ["VITE_API_URL", "Explicit API base (e.g. https://plant.example.com/api); empty = same-origin /api"],
            ["VITE_GOOGLE_CLIENT_ID", "Google OAuth button (GoogleOAuthProvider)"],
            ["VITE_AZURE_CLIENT_ID / VITE_AZURE_TENANT_ID", "MSAL configuration (msalConfig.js)"],
        ],
    )
    doc.add_heading("12.3 Startup Sequence", level=2)
    add_numbered(
        doc,
        [
            "Create MySQL database; run cd backend && npm run db:schema",
            "Apply migrate_*.sql as needed; npm run db:migrate:deploy for Prisma history",
            "Configure backend/.env (JWT_SECRET, DATABASE_URL, SMTP, CLIENT_ORIGIN)",
            "cd backend && npm run dev (or npm start for production)",
            "Configure frontend/.env; cd frontend && npm run dev (or npm run build + serve dist)",
        ],
    )
    doc.add_page_break()

    # ── 13. Security ──
    doc.add_heading("13. Security", level=1)
    add_table(
        doc,
        ["Control", "Implementation"],
        [
            ["Transport", "HTTPS expected in production; HSTS via helmet when trust proxy + HTTPS"],
            ["Authentication", "JWT Bearer; bcrypt password hashing; SSO token verification"],
            ["Brute force", "loginLimiter 30 attempts per 15 minutes on /api/auth/login"],
            ["Rate limiting", "globalLimiter 500 req/15min per IP"],
            ["CORS", "Restricted to CLIENT_ORIGIN with credentials"],
            ["Headers", "helmet (CSP disabled for JSON API; CORP cross-origin for avatars)"],
            ["Authorization", "Role-based admin routes; mapping-based form access"],
            ["File upload", "MIME/size limits; multer disk storage; authenticated download routes"],
            ["SQL injection", "Parameterized queries throughout controllers"],
            ["Secrets", "JWT_SECRET min 32 chars; .env not committed (use .env.example)"],
            ["Avatar access", "GET /api/auth/users/:userId/avatar requires Bearer token"],
        ],
    )
    doc.add_page_break()

    # ── 14. Performance ──
    doc.add_heading("14. Performance", level=1)
    add_table(
        doc,
        ["Area", "Tuning"],
        [
            ["BI date bounds", "DEFAULT_LOOKBACK_DAYS=365 prevents full-table scans on dashboard load"],
            ["BI row cap", "BI_ROW_LIMIT=200000 hard limit on time-series query results"],
            ["Connection pool", "DB_POOL_LIMIT=30 (keep below MySQL max_connections)"],
            ["Pool queue", "DB_POOL_QUEUE_LIMIT=100 — excess requests fail fast"],
            ["Keep-alive", "enableKeepAlive=true, keepAliveInitialDelay=10000 on mysql2 pool"],
            ["Date strings", "dateStrings:true avoids timezone JSON serialization issues"],
            ["CSV export", "getRecords allows limit up to 10000 per request"],
            ["JSON body", "10mb limit for large equipment image metadata payloads"],
        ],
    )
    doc.add_page_break()

    # ── 15. Known Limitations ──
    doc.add_heading("15. Known Limitations", level=1)
    add_table(
        doc,
        ["Limitation", "Technical Impact"],
        [
            ["No approval workflow", "FormReviewModal is self-certification only"],
            ["No logbook edit/delete API", "Committed form rows are immutable via UI/API"],
            ["No submitter attribution", "Logbook tables lack user_id column"],
            ["Batch submit no duplicate check", "prod pan/decanter/clarification can insert duplicates"],
            ["ehs_accident not routed", "Table + FORM_CONFIG exist; App.jsx has no live route"],
            ["Equipment access unscoped", "Any authenticated user can edit equipment history"],
            ["Prisma not used at runtime", "Schema changes need manual sync between init.sql, migrate_*.sql, Prisma"],
            ["No Docker", "Deployment procedures are environment-specific"],
            ["No structured audit log", "Only stderr console.error via logServerError"],
            ["Manager field unused", "manager_id not linked to workflows"],
        ],
    )
    doc.add_page_break()

    # ── 16. Troubleshooting ──
    doc.add_heading("16. Troubleshooting", level=1)
    add_table(
        doc,
        ["Symptom", "Likely Cause", "Resolution"],
        [
            ["FATAL: JWT_SECRET must be set", "Missing/short JWT_SECRET in .env", "Set JWT_SECRET ≥ 32 characters"],
            ["MySQL connection failed on startup", "Wrong DATABASE_URL or MySQL not running", "Verify credentials; check MYSQL_HOST/PORT"],
            ["401 Invalid or expired token", "JWT expired or JWT_SECRET changed", "Re-login; ensure consistent JWT_SECRET across restarts"],
            ["403 Access denied to this form", "Missing mapping or mapping_forms restriction", "Admin: assign app/form in Employee Management"],
            ["403 SSO access denied", "Email not pre-provisioned in users", "Admin creates user before first SSO login"],
            ["409 Duplicate operation", "Same date/shift/time already exists", "Change operational key fields or request DB correction"],
            ["CORS error in browser", "CLIENT_ORIGIN mismatch", "Set CLIENT_ORIGIN to exact frontend URL"],
            ["429 Too many requests", "Rate limit exceeded", "Wait 15 minutes or adjust limiters for dev"],
            ["BI dashboard empty", "No data in date range", "Check DEFAULT_LOOKBACK_DAYS window; verify form submissions exist"],
            ["Avatar not loading", "Missing token on image request", "Use AuthenticatedImage component; check Bearer header"],
            ["Pool queue timeout", "DB_POOL_LIMIT exhausted", "Increase pool or reduce concurrent BI queries"],
        ],
    )
    doc.add_page_break()

    # ── 17. Appendix ──
    doc.add_heading("17. Appendix", level=1)
    doc.add_heading("17.1 Backend npm Scripts (package.json)", level=2)
    add_table(
        doc,
        ["Script", "Command"],
        [
            ["start", "node server.js"],
            ["dev", "nodemon server.js"],
            ["db:schema", "node scripts/apply-init-sql.js"],
            ["db:apply-sql", "node scripts/apply-sql-file.js"],
            ["db:migrate:dev", "prisma migrate dev"],
            ["db:migrate:deploy", "prisma migrate deploy"],
            ["db:migrate:status", "prisma migrate status"],
            ["db:migrate:resolve-baseline", "prisma migrate resolve --applied 20260213120000_baseline"],
            ["db:generate", "prisma generate"],
            ["db:seed-ppn-hierarchy", "node scripts/seed-ppn-hierarchy.js"],
            ["db:mill-mapping", "node scripts/import-mill-mapping.js"],
            ["db:add-mill-house-equipment", "node scripts/add-mill-house-equipment.js"],
            ["db:import-power-xlsx", "node scripts/import-power-equipment-xlsx.js"],
            ["db:import-ppn-feed-all", "node scripts/data_feed_power_history/import-all-feeds.js"],
            ["db:clear-power-equipment", "node scripts/clear-power-equipment-data.js"],
            ["backlog:import-lab", "node scripts/import-backlog-lab-csv.js"],
            ["seed", "node seed.js"],
        ],
    )
    doc.add_heading("17.2 Frontend npm Scripts", level=2)
    add_table(
        doc,
        ["Script", "Command"],
        [
            ["dev", "vite (port 5173)"],
            ["build", "vite build → dist/"],
            ["preview", "vite preview"],
        ],
    )
    doc.add_heading("17.3 Pattern Reference Quick Guide", level=2)
    add_table(doc, ["Pattern", "Columns", "Example Forms"], PATTERN_DESCRIPTIONS)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("— End of Document —").italic = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    DOCS_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUTPUT_DOCX))
    return OUTPUT_DOCX


def build_markdown() -> Path:
    md = MdBuilder()

    md.h(1, "DigiLog Technical Documentation")
    md.p(f"**Zuari Industries — Digital Operations Platform**")
    md.p(f"Document Version: 1.0  ")
    md.p(f"Date: {DOC_DATE}  ")
    md.p("Audience: Developers, DevOps Engineers, Security Auditors, System Maintainers")
    md.p("---")

    md.h(2, "Document Control")
    md.table(
        ["Item", "Details"],
        [
            ["Document Title", "DigiLog Technical Documentation"],
            ["System Name", "DigiLog — Digital Logbook Platform"],
            ["Organization", "Zuari Industries"],
            ["Purpose", "Technical reference for architecture, APIs, database, deployment, and maintenance"],
            ["Classification", "Internal — Technical Use"],
            ["Repository", str(ROOT)],
        ],
    )

    # 1
    md.h(2, "1. Project Overview")
    md.p(
        "DigiLog is Zuari Industries' plant-wide digital logbook and operations platform. "
        "It replaces paper registers with a React single-page application backed by an Express REST API "
        "and MySQL 8 database."
    )
    md.h(3, "1.1 Repository Layout")
    md.bullets(
        [
            "`frontend/` — React 18 SPA (Vite dev server port 5173)",
            "`backend/` — Express API (default port 5000)",
            "`mysql/` — init.sql baseline schema and migrate_*.sql patches",
            "`scripts/` — Python/Node utilities including this documentation generator",
            "`docs/` — Generated documentation outputs",
        ]
    )
    md.h(3, "1.2 Key Technical Characteristics")
    md.table(
        ["Characteristic", "Implementation"],
        [
            ["API style", "REST JSON under `/api/*` prefix"],
            ["Auth", "JWT Bearer tokens (7d default); roles `admin` | `employee`"],
            ["Data access", "mysql2 raw SQL at runtime; Prisma for migrations only"],
            ["Access control", "mappings, mapping_forms, user_homepage_cards, user_data_upload_access"],
            ["Form workflow", "FormReviewModal = pre-submit self-certification only"],
            ["Deployment", "No Docker in repo; Node + MySQL + static Vite build"],
        ],
    )

    # 2
    md.h(2, "2. System Architecture")
    md.h(3, "2.1 Three-Tier Architecture (Mermaid)")
    md.code(ARCH_MERMAID, "mermaid")
    md.h(3, "2.2 Request Flow — Form Submit")
    md.numbered(
        [
            "Employee opens `/forms/<formKey>`.",
            "Submit → `FormReviewModal` read-only summary.",
            "Confirm & Commit → `POST /api/forms/:formKey`.",
            "Axios attaches `Authorization: Bearer <JWT>`.",
            "`authenticate` middleware validates token.",
            "`canAccessForm` checks mappings/mapping_forms.",
            "`injectDateCols` applies pattern A–H columns.",
            "`hasDuplicateOperationRow` blocks duplicates (single submit).",
            "INSERT into target table; 201 response.",
        ]
    )
    md.h(3, "2.3 Portal Separation")
    md.p(
        "`enforceAdminPortalRules` in auth.controller.js blocks cross-portal login. "
        "SSO users must be pre-provisioned in the `users` table."
    )

    # 3
    md.h(2, "3. Technology Stack")
    md.table(
        ["Layer", "Technology", "Version / Notes"],
        [
            ["Frontend", "React + Vite + Tailwind", "18.3 / 6.4 / 3.4"],
            ["HTTP / Auth UI", "Axios + MSAL + Google OAuth", "1.7 / 3.14 / 0.12"],
            ["Charts", "Recharts", "3.8"],
            ["Backend", "Express + mysql2", "4.19 / 3.10"],
            ["Migrations", "Prisma", "7.8 (@@ignore models)"],
            ["Security", "helmet + rate-limit + bcrypt + JWT", "8.2 / 7.3"],
            ["Upload / Email", "multer + nodemailer", "2.1 / 6.9"],
            ["Database", "MySQL", "8.x utf8mb4"],
        ],
    )

    # 4
    md.h(2, "4. Folder Structure")
    md.table(
        ["Path", "Purpose"],
        [
            ["`frontend/src/App.jsx`", "All React routes"],
            ["`frontend/src/pages/`", "Page components"],
            ["`frontend/src/components/`", "Shared UI including FormReviewModal"],
            ["`frontend/src/hooks/`", "useAuth, useFormMeta, usePowerPlantHierarchy, etc."],
            ["`frontend/src/context/AuthContext.jsx`", "JWT + MSAL + Google session"],
            ["`backend/server.js`", "Express bootstrap"],
            ["`backend/routes/`", "10 route modules"],
            ["`backend/controllers/`", "Business logic"],
            ["`mysql/init.sql`", "Baseline schema (~49 tables)"],
            ["`mysql/migrate_*.sql`", "23 incremental SQL patches"],
        ],
    )

    # 5
    md.h(2, "5. Frontend Documentation")
    md.h(3, "5.1 Routes (App.jsx)")
    md.p(
        "Public: `/`, `/admin/login`, `/operations-desk`. Protected: `/dashboard`, `/forms-hub`, `/bi/*`, "
        "`/forms/*` (24 form keys), `/equipment/*`, `/power/*`, `/power-plant-equipment-new/*`, "
        "`/ehs/*`, `/production/*`, `/data-upload`. Admin-only: `/admin/employees`."
    )
    md.h(3, "5.2 Components")
    md.bullets(
        [
            "`ProtectedRoute` — auth + optional `requiredRole='admin'`",
            "`FormReviewModal` — pre-submit review (not approval workflow)",
            "`FormTable` — View Data modal + CSV export (limit 10000)",
            "`PowerPlantHierarchyExplorer` — PPN tree UI",
            "`DistilleryChartsGrid` / `MillRawDataTable` — BI visualizations",
        ]
    )
    md.h(3, "5.3 Hooks & Context")
    md.table(
        ["Module", "Purpose"],
        [
            ["`AuthContext.jsx`", "JWT, loginManual, MSAL redirect, Google login, logout"],
            ["`useAuth.js`", "Auth context consumer"],
            ["`useGsmaFormReview.js`", "Submit → review modal → API POST"],
            ["`useFormMeta.js`", "Form metadata from API"],
            ["`usePowerPlantHierarchy.js`", "PPN hierarchy tree"],
            ["`useDataUploadAccess.js`", "Upload permission gate"],
        ],
    )

    # 6
    md.h(2, "6. Backend Documentation")
    md.h(3, "6.1 server.js")
    md.bullets(
        [
            "helmet, cors, json 10mb, global + login rate limiters",
            "Mounts `/api/auth`, `/api/admin`, `/api/apps`, `/api/forms`, `/api/equipment`, `/api/power`, `/api/power-new`, `/api/bi`, `/api/homepage-cards`, `/api/data-upload`",
            "Global error handler with `mapDbError`",
        ]
    )
    md.h(3, "6.2 Controllers")
    md.p(
        "auth, admin, form (FORM_CONFIG), app, equipment, power, powerNew, ppnHierarchy, "
        "bi, biSettings, homepageCards, dataUpload."
    )

    # 7
    md.h(2, "7. API Documentation")
    md.p("All endpoints under `/api/*`. Full listing:")
    md.table(
        ["Method", "Endpoint", "Auth", "Description"],
        API_ENDPOINTS,
    )

    # 8
    md.h(2, "8. Database Documentation")
    md.h(3, "8.1 Schema Management")
    md.bullets(
        [
            "Runtime: mysql2 raw SQL (Prisma Client not used)",
            "Baseline: `mysql/init.sql` via `npm run db:schema`",
            "Patches: 23 `migrate_*.sql` files",
            "Prisma migrations track form/logbook evolution",
            "Schema drift possible between init.sql, migrate_*.sql, and Prisma — verify before deploy",
        ]
    )
    md.h(3, "8.2 Tables (49)")
    md.table(["Table", "Domain", "Description"], DATABASE_TABLES)
    md.h(3, "8.3 Access Control Tables")
    md.table(
        ["Table", "Relationship"],
        [
            ["`mappings`", "user_id → app_id"],
            ["`mapping_forms`", "Optional form_id restriction per mapping"],
            ["`user_homepage_cards`", "Dashboard card visibility"],
            ["`user_data_upload_access`", "Data Ingestion Center grant"],
        ],
    )
    schemas = load_table_schemas()
    render_schema_sections_md(md, schemas, table_description_map())

    # 9
    md.h(2, "9. Authentication & Authorization")
    md.h(3, "9.1 JWT")
    md.p("Token in localStorage; Axios Bearer interceptor; 401 → redirect `/?login=1`.")
    md.h(3, "9.2 Roles")
    md.table(
        ["Role", "Portal", "Notes"],
        [
            ["employee", "Home `/`", "Mapping-gated forms and BI"],
            ["admin", "`/admin/login`", "Employee management; blocked from employee portal"],
        ],
    )
    md.h(3, "9.3 SSO")
    md.p("Microsoft/Google require pre-created `users` row; first login links provider ID.")

    # 10
    md.h(2, "10. Business Logic")
    md.h(3, "10.1 FORM_CONFIG")
    md.table(["Pattern", "Key Columns", "Used By"], PATTERN_DESCRIPTIONS)
    md.table(["formKey", "Table", "Pattern", "Duplicate Key", "Notes"], FORM_CONFIG_ROWS)
    md.h(3, "10.2 Duplicate vs Batch")
    md.p(
        "Single submit: duplicate detection via `hasDuplicateOperationRow`. "
        "Batch (`POST .../batch`): no duplicate check — `prod_pan_logbook`, `prod_decanter`, `prod_clarification`."
    )
    md.h(3, "10.3 FormReviewModal")
    md.p("Pre-submit self-certification only. No formal approval workflow.")

    # 11
    md.h(2, "11. Logging & Audit")
    md.table(
        ["Mechanism", "Logged", "Gap"],
        [
            ["logServerError", "stderr stack traces", "No centralized logging"],
            ["Logbook timestamps", "INSERT timestamp column", "No user_id on rows"],
            ["data_upload_files", "Uploader audit", "—"],
            ["Equipment history", "created_at/updated_at", "No change history"],
        ],
    )

    # 12
    md.h(2, "12. Deployment & Environment Setup")
    md.p("No Docker in repo. Typical: MySQL 8 + Node backend + nginx serving Vite `dist/`.")
    md.h(3, "12.1 Backend env.js")
    md.table(
        ["Variable", "Default", "Purpose"],
        [
            ["PORT", "5000", "Express port"],
            ["JWT_SECRET", "required ≥32", "Token signing"],
            ["JWT_EXPIRES_IN", "7d", "Token TTL"],
            ["DATABASE_URL", "from MYSQL_*", "MySQL connection"],
            ["CLIENT_ORIGIN", "http://localhost:5173", "CORS"],
            ["SMTP_*", "empty", "Activation emails"],
            ["DATA_UPLOAD_MAX_BYTES", "26214400", "25 MB upload cap"],
            ["DB_POOL_LIMIT", "30", "Connection pool size"],
            ["DB_POOL_QUEUE_LIMIT", "100", "Queue bound"],
            ["DB_CONNECT_TIMEOUT", "10000", "Connect timeout ms"],
        ],
    )
    md.h(3, "12.2 Frontend")
    md.bullets(
        [
            "`VITE_API_URL` — API base URL",
            "`VITE_GOOGLE_CLIENT_ID` — Google OAuth",
            "`VITE_AZURE_CLIENT_ID` / `VITE_AZURE_TENANT_ID` — MSAL",
        ]
    )

    # 13
    md.h(2, "13. Security")
    md.bullets(
        [
            "JWT Bearer + bcrypt; SSO token verification",
            "Rate limits: 30 login / 500 global per 15 min",
            "CORS restricted to CLIENT_ORIGIN",
            "helmet headers; parameterized SQL",
            "Authenticated avatar and file download routes",
        ]
    )

    # 14
    md.h(2, "14. Performance")
    md.bullets(
        [
            "BI `DEFAULT_LOOKBACK_DAYS=365` date bound on dashboard load",
            "BI `BI_ROW_LIMIT=200000` query cap",
            "Pool: `DB_POOL_LIMIT=30`, `DB_POOL_QUEUE_LIMIT=100`, keep-alive enabled",
            "CSV export: up to 10000 records per `getRecords` request",
        ]
    )

    # 15
    md.h(2, "15. Known Limitations")
    md.table(
        ["Limitation", "Impact"],
        [
            ["No approval workflow", "FormReviewModal only"],
            ["Immutable logbook rows", "No edit/delete API"],
            ["No submitter on logbooks", "Audit gap"],
            ["Batch no duplicate check", "Possible duplicate production rows"],
            ["ehs_accident unrouted", "Table exists, no App.jsx route"],
            ["Schema drift risk", "init.sql vs migrate vs Prisma"],
        ],
    )

    # 16
    md.h(2, "16. Troubleshooting")
    md.table(
        ["Symptom", "Cause", "Fix"],
        [
            ["JWT_SECRET FATAL", "Missing secret", "Set ≥32 char JWT_SECRET"],
            ["MySQL connection failed", "Bad DATABASE_URL", "Check MySQL running and credentials"],
            ["403 form access", "No mapping", "Admin assigns mapping"],
            ["403 SSO denied", "User not provisioned", "Admin creates user first"],
            ["409 duplicate", "Same operational key", "Change date/shift/time"],
            ["CORS error", "CLIENT_ORIGIN mismatch", "Align env with frontend URL"],
        ],
    )

    # 17
    md.h(2, "17. Appendix")
    md.h(3, "17.1 Backend npm Scripts")
    md.table(
        ["Script", "Command"],
        [
            ["start", "node server.js"],
            ["dev", "nodemon server.js"],
            ["db:schema", "apply init.sql"],
            ["db:migrate:deploy", "prisma migrate deploy"],
            ["db:seed-ppn-hierarchy", "seed hierarchy nodes"],
            ["backlog:import-lab", "import lab CSV backlog"],
        ],
    )
    md.h(3, "17.2 Frontend npm Scripts")
    md.table(
        ["Script", "Command"],
        [
            ["dev", "vite"],
            ["build", "vite build"],
            ["preview", "vite preview"],
        ],
    )

    md.p("---")
    md.p("*— End of Document —*")

    DOCS_DIR.mkdir(parents=True, exist_ok=True)
    OUTPUT_MD.write_text(md.build(), encoding="utf-8")
    return OUTPUT_MD


def main():
    docx_path = build_docx()
    md_path = build_markdown()
    docx_size = docx_path.stat().st_size
    md_size = md_path.stat().st_size
    print(f"Generated: {docx_path}")
    print(f"  Size: {docx_size:,} bytes ({docx_size / 1024:.1f} KB)")
    print(f"Generated: {md_path}")
    print(f"  Size: {md_size:,} bytes ({md_size / 1024:.1f} KB)")


if __name__ == "__main__":
    main()
